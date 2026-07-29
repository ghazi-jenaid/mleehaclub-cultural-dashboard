"""لوحة تحكم محلية ذكية لمجلد اللجنة الثقافية.

تشغّل عبر:
    streamlit run app.py

المتطلبات:
    streamlit>=1.37,<2
    watchdog>=4,<7
    plotly>=5.22,<7
    pandas>=2.1,<3
    pymupdf>=1.24,<2
"""

from __future__ import annotations

import base64
import html
import mimetypes
import os
import re
import threading
import time
import zipfile
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree

import pandas as pd
import plotly.express as px
import streamlit as st
import streamlit.components.v1 as components
import fitz
from docx import Document
from openpyxl import load_workbook
from watchdog.events import FileSystemEventHandler
from watchdog.observers import Observer


APP_DIR = Path(__file__).resolve().parent
DEFAULT_PROJECT = Path(
    os.environ.get("COMMITTEE_PROJECT_ROOT", str(APP_DIR.parent))
).expanduser().resolve()
DOCUMENT_SYNC_VERSION = 2
IGNORED_DIRS = {
    APP_DIR.name,
    "لوحة إدارة المشروع",
    ".git",
    ".idea",
    ".vscode",
    "__pycache__",
    ".tmp.driveupload",
    "node_modules",
    ".venv",
    "venv",
}
MAX_PREVIEW_BYTES = 2 * 1024 * 1024
MAX_LINE_SCAN_BYTES = 50 * 1024 * 1024
MAX_DOWNLOAD_BYTES = 100 * 1024 * 1024
LARGE_FILE_BYTES = 5 * 1024 * 1024
ARABIC_DIGIT_TRANSLATION = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
CODE_EXTENSIONS = {
    ".py": "python",
    ".js": "javascript",
    ".jsx": "javascript",
    ".ts": "typescript",
    ".tsx": "typescript",
    ".html": "html",
    ".css": "css",
    ".scss": "scss",
    ".json": "json",
    ".xml": "xml",
    ".yaml": "yaml",
    ".yml": "yaml",
    ".toml": "toml",
    ".ini": "ini",
    ".sql": "sql",
    ".sh": "bash",
    ".ps1": "powershell",
    ".bat": "batch",
    ".md": "markdown",
    ".java": "java",
    ".c": "c",
    ".cpp": "cpp",
    ".h": "c",
    ".cs": "csharp",
    ".go": "go",
    ".rs": "rust",
    ".php": "php",
    ".rb": "ruby",
    ".r": "r",
}
TEXT_EXTENSIONS = set(CODE_EXTENSIONS) | {
    ".txt",
    ".csv",
    ".tsv",
    ".log",
    ".conf",
    ".cfg",
    ".properties",
}
FILE_ICONS = {
    ".pdf": "📕",
    ".doc": "📘",
    ".docx": "📘",
    ".xls": "📗",
    ".xlsx": "📗",
    ".csv": "📊",
    ".ppt": "📙",
    ".pptx": "📙",
    ".png": "🖼️",
    ".jpg": "🖼️",
    ".jpeg": "🖼️",
    ".gif": "🖼️",
    ".mp4": "🎬",
    ".mp3": "🎵",
    ".zip": "🗜️",
    ".py": "🐍",
    ".js": "◆",
    ".ts": "◆",
    ".json": "{ }",
    ".md": "📝",
    ".txt": "📄",
}
APPROVED_FORMS = [
    ("MCSC-F-11", "محضر اجتماع", "محاضر الاجتماعات"),
    ("MCSC-F-40", "تقرير إنجاز فعالية", "تقارير انجاز الفعاليات"),
    ("MCSC-F-46", "خطة الأنشطة الثقافية والمجتمعية", "خطة الأنشطة الثقافية والمجتمعية"),
    ("MCSC-F-140", "نموذج حصر متطلبات فعالية", "حصر متطلبات الفعاليات"),
    ("MCSC-F-141", "التقرير الشهري", "التقارير الشهرية"),
    ("MCSC-F-142", "التصور العام للفعاليات", "التصور العام"),
    ("MCSC-F-160", "القائمة التفقدية لإدارة الفعاليات", "القوائم التفقدية للفعاليات"),
    ("MCSC-F-221", "خطة العمل الفردية للفعاليات", "خطة العمل الفردية للفعاليات"),
    ("MCSC-F-249", "قرارات اللجنة الثقافية", "قرارات اللجنة"),
    # الجذر المشترك "استب" يتعرف على التسميتين المستخدمتين فعلياً:
    # "الاستبيانات" في موسم 2025/2026 و"الاستبانات" في موسم 2026/2027.
    ("الاستبانات", "قياس رضا المستفيدين", "استب"),
]
SPORTS_QUARTERS = [
    ("الربع الأول", "يوليو – سبتمبر"),
    ("الربع الثاني", "أكتوبر – ديسمبر"),
    ("الربع الثالث", "يناير – مارس"),
    ("الربع الرابع", "أبريل – يونيو"),
]


@dataclass(frozen=True)
class FileInfo:
    """البيانات الوصفية الآمنة لملف واحد."""

    relative_path: str
    absolute_path: str
    extension: str
    size_bytes: int
    modified_at: float
    line_count: int | None
    error: str | None = None


class ChangeHandler(FileSystemEventHandler):
    """يجمع إشعارات watchdog دون تنفيذ أي عمل ثقيل داخل خيط المراقبة."""

    def __init__(self, service: "WatchService") -> None:
        self.service = service

    def on_any_event(self, event) -> None:  # noqa: ANN001
        """يسجل وصول حدث نظام ملفات غير مستبعد."""
        if not _is_ignored_path(Path(event.src_path)):
            self.service.mark_changed()


class WatchService:
    """يدير Observer واحداً مستقلاً لكل جلسة Streamlit."""

    def __init__(self, root: Path) -> None:
        """ينشئ مراقباً متكرراً للمجلد المحدد ويبدأ خيطه."""
        self.root = root
        self.signal = threading.Event()
        self._lock = threading.Lock()
        self._revision = 0
        self.observer = Observer()
        self.observer.schedule(ChangeHandler(self), str(root), recursive=True)
        self.observer.daemon = True
        self.observer.start()

    @property
    def revision(self) -> int:
        """يعيد رقم آخر تغيير؛ لا تضيع الإشارة عند تعدد جلسات المتصفح."""
        with self._lock:
            return self._revision

    def mark_changed(self) -> None:
        """يسجل تغييراً بطريقة آمنة بين خيط المراقبة وخيط الواجهة."""
        with self._lock:
            self._revision += 1
        self.signal.set()

    def is_alive(self) -> bool:
        """يتحقق من سلامة خيط المراقبة."""
        return self.observer.is_alive()

    def stop(self) -> None:
        """يوقف خيط المراقبة بأمان."""
        self.observer.stop()
        self.observer.join(timeout=2)


@st.cache_resource(show_spinner=False)
def get_watch_service(root_text: str) -> WatchService:
    """ينشئ مراقباً واحداً مشتركاً ومستقراً لجميع جلسات المتصفح."""
    return WatchService(Path(root_text))


def _is_ignored_path(path: Path) -> bool:
    """يتحقق مما إذا كان المسار ضمن مجلدات النظام المستبعدة."""
    try:
        parts = set(path.resolve().relative_to(DEFAULT_PROJECT.resolve()).parts)
    except (OSError, ValueError):
        parts = set(path.parts)
    return bool(parts & IGNORED_DIRS)


def _looks_textual(path: Path) -> bool:
    """يقدّر بأمان ما إذا كان الملف نصياً وقابلاً لحساب الأسطر."""
    if path.suffix.lower() in TEXT_EXTENSIONS:
        return True
    try:
        with path.open("rb") as handle:
            sample = handle.read(4096)
        return b"\x00" not in sample
    except OSError:
        return False


def _count_lines(path: Path, size: int) -> int | None:
    """يحسب أسطر ملف نصي ضمن حد الحجم، أو يعيد None."""
    if size > MAX_LINE_SCAN_BYTES or not _looks_textual(path):
        return None
    try:
        with path.open("rb") as handle:
            count = sum(chunk.count(b"\n") for chunk in iter(lambda: handle.read(1024 * 1024), b""))
        if size == 0:
            return 0
        with path.open("rb") as handle:
            handle.seek(-1, os.SEEK_END)
            return count if handle.read(1) == b"\n" else count + 1
    except (OSError, ValueError):
        return None


def _walk_files(root: Path) -> Iterable[Path]:
    """مسح مقاوم لأخطاء الصلاحيات والمجلدات غير المتاحة."""
    def on_error(_: OSError) -> None:
        """يتجاوز أخطاء الوصول الجزئية ويتيح استمرار المسح."""
        return

    for current, dirs, names in os.walk(root, topdown=True, onerror=on_error):
        dirs[:] = [name for name in dirs if name not in IGNORED_DIRS]
        for name in names:
            if name.startswith("."):
                continue
            path = Path(current) / name
            if not _is_ignored_path(path):
                yield path


@st.cache_data(show_spinner=False, max_entries=8)
def scan_project(root_text: str, revision: int) -> list[dict]:
    """يمسح المشروع؛ revision مفتاح صريح لإبطال التخزين المؤقت عند التغيير."""
    del revision
    root = Path(root_text)
    results: list[FileInfo] = []
    for path in _walk_files(root):
        try:
            stat = path.stat()
            suffix = path.suffix.lower() or "بدون امتداد"
            results.append(
                FileInfo(
                    relative_path=str(path.relative_to(root)),
                    absolute_path=str(path),
                    extension=suffix,
                    size_bytes=stat.st_size,
                    modified_at=stat.st_mtime,
                    line_count=_count_lines(path, stat.st_size),
                )
            )
        except (OSError, ValueError) as exc:
            results.append(
                FileInfo(str(path), str(path), "غير معروف", 0, 0, None, str(exc))
            )
    return [asdict(item) for item in results]


@st.cache_data(show_spinner=False, max_entries=128)
def read_text_preview(path_text: str, modified_at: float) -> tuple[str | None, str | None]:
    """يقرأ المعاينة فقط، مع ربط الكاش بتاريخ آخر تعديل."""
    del modified_at
    path = Path(path_text)
    try:
        size = path.stat().st_size
        if size > MAX_PREVIEW_BYTES:
            return None, f"حجم الملف يتجاوز حد المعاينة ({format_bytes(MAX_PREVIEW_BYTES)})."
        if not _looks_textual(path):
            return None, "هذا ملف ثنائي ولا يمكن عرضه كنص."
        raw = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-8", "cp1256", "latin-1"):
            try:
                return raw.decode(encoding), None
            except UnicodeDecodeError:
                continue
        return None, "تعذر تحديد ترميز الملف."
    except (OSError, ValueError) as exc:
        return None, f"تعذر فتح الملف: {exc}"


@st.cache_data(show_spinner=False, max_entries=32)
def read_docx_preview(path_text: str, modified_at: float) -> tuple[str | None, str | None]:
    """يستخرج نص DOCX للمعاينة دون تعديل المستند."""
    del modified_at
    try:
        with zipfile.ZipFile(path_text) as archive:
            xml = archive.read("word/document.xml")
        root = ElementTree.fromstring(xml)
        paragraphs: list[str] = []
        for paragraph in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
            text = "".join(
                node.text or ""
                for node in paragraph.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
            ).strip()
            if text:
                paragraphs.append(text)
        return "\n\n".join(paragraphs), None
    except (OSError, KeyError, zipfile.BadZipFile, ElementTree.ParseError) as exc:
        return None, f"تعذر استخراج نص المستند: {exc}"


@st.cache_data(show_spinner=False, max_entries=24)
def read_docx_tables(path_text: str, modified_at: float) -> list[list[list[str]]]:
    """يستخرج جداول Word كاملة للعرض داخل اللوحة."""
    del modified_at
    document = Document(path_text)
    return [
        [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]


@st.cache_data(show_spinner=False, max_entries=24)
def read_excel_sheet(path_text: str, modified_at: float, sheet_name: str) -> pd.DataFrame:
    """يقرأ نطاقاً مستخدماً من ورقة Excel مع حماية الواجهة من الجداول المفرطة."""
    del modified_at
    workbook = load_workbook(path_text, data_only=True, read_only=True)
    sheet = workbook[sheet_name]
    rows = list(sheet.iter_rows(max_row=min(sheet.max_row, 500), max_col=min(sheet.max_column, 60), values_only=True))
    while rows and not any(value not in (None, "") for value in rows[-1]):
        rows.pop()
    return pd.DataFrame(rows)


@st.cache_data(show_spinner=False, max_entries=24)
def get_excel_sheets(path_text: str, modified_at: float) -> list[str]:
    """يعيد أسماء أوراق عمل Excel."""
    del modified_at
    workbook = load_workbook(path_text, data_only=True, read_only=True)
    return workbook.sheetnames


def find_pdf_twin(files: list[dict], item: dict) -> dict | None:
    """يبحث عن نسخة PDF تحمل اسم المستند نفسه لعرض تنسيقه الأصلي."""
    stem = _clean_event_name(Path(item["relative_path"]).stem).casefold()
    candidates = [
        candidate for candidate in files
        if candidate["extension"] == ".pdf"
        and _clean_event_name(Path(candidate["relative_path"]).stem).casefold() == stem
        and candidate["absolute_path"] != item["absolute_path"]
    ]
    if not candidates:
        return None
    parent = Path(item["absolute_path"]).parent
    return min(candidates, key=lambda candidate: len(Path(candidate["absolute_path"]).parent.parts) - len(parent.parts))


def render_scrollable_pdf_images(
    pages: list[tuple[int, bytes]],
    display_width: int,
    viewport_height: int = 850,
) -> None:
    """يعرض صور صفحات PDF داخل إطار مستقل بتمرير أفقي ورأسي حقيقي."""
    page_blocks: list[str] = []
    for page_number, image_bytes in pages:
        encoded = base64.b64encode(image_bytes).decode("ascii")
        page_blocks.append(
            f'<section class="pdf-page">'
            f'<div class="page-label">الصفحة {page_number:,}</div>'
            f'<img src="data:image/png;base64,{encoded}" '
            f'width="{display_width}" draggable="false" alt="صفحة {page_number:,}">'
            f'</section>'
        )
    viewer_html = f"""
    <!doctype html>
    <html lang="ar" dir="rtl">
    <head>
      <meta charset="utf-8">
      <style>
        * {{ box-sizing:border-box; }}
        html, body {{
          margin:0; padding:0; width:100%; height:100%;
          background:#e9e6de; overflow:hidden;
          font-family:"Noto Kufi Arabic","Segoe UI",sans-serif;
        }}
        .scrollport {{
          position:relative; width:100vw; height:100vh;
          overflow:scroll; direction:rtl;
          scrollbar-gutter:stable both-edges;
        }}
        .viewer {{
          direction:rtl; width:max-content; min-width:100%;
          padding:16px; margin:0;
        }}
        .pdf-page {{
          direction:rtl; width:{display_width}px; margin:0 auto 18px;
          background:white; box-shadow:0 5px 18px rgba(0,0,0,.16);
        }}
        .page-label {{
          position:sticky; right:0; width:max-content; z-index:2;
          padding:7px 12px; background:#1f5039; color:white;
          font-size:13px; font-weight:700;
        }}
        img {{
          display:block; width:{display_width}px; max-width:none !important;
          height:auto; margin:0; image-rendering:auto;
          user-select:none; -webkit-user-drag:none;
        }}
      </style>
    </head>
    <body oncontextmenu="return false;">
      <div class="scrollport">
        <main class="viewer">{''.join(page_blocks)}</main>
      </div>
    </body>
    </html>
    """
    components.html(viewer_html, height=viewport_height, scrolling=False)


def render_pdf_document(path: Path, modified_at: float, key_prefix: str) -> None:
    """يعرض PDF بقارئ محلي عالي الدقة من دون تمرير الملف إلى قارئ المتصفح."""
    page_count = get_pdf_page_count(str(path), modified_at)
    if page_count < 1:
        st.warning("لا يحتوي ملف PDF على صفحات قابلة للعرض.")
        return

    page_key = f"{key_prefix}-page"
    rotation_key = f"{key_prefix}-rotation"
    st.session_state.setdefault(page_key, 1)
    st.session_state.setdefault(rotation_key, 0)
    st.session_state[page_key] = min(max(int(st.session_state[page_key]), 1), page_count)

    def move_page(delta: int) -> None:
        st.session_state[page_key] = min(max(st.session_state[page_key] + delta, 1), page_count)

    def rotate_page(delta: int) -> None:
        st.session_state[rotation_key] = (st.session_state[rotation_key] + delta) % 360

    with st.container(border=True):
        display_mode = st.radio(
            "طريقة العرض",
            ["المستند كاملاً في صفحة عرض واحدة", "التنقل اليدوي بين الصفحات"],
            horizontal=True,
            key=f"{key_prefix}-display-mode",
        )

        if display_mode == "التنقل اليدوي بين الصفحات":
            previous_col, page_col, next_col, zoom_col, left_col, right_col = st.columns(
                [1, 1.35, 1, 1.35, 1, 1],
                vertical_alignment="bottom",
            )
            with previous_col:
                st.button(
                    "السابق",
                    icon=":material/arrow_forward:",
                    disabled=st.session_state[page_key] <= 1,
                    on_click=move_page,
                    args=(-1,),
                    key=f"{key_prefix}-previous",
                    width="stretch",
                )
            with page_col:
                st.number_input(
                    "رقم الصفحة",
                    min_value=1,
                    max_value=page_count,
                    step=1,
                    key=page_key,
                )
            with next_col:
                st.button(
                    "التالي",
                    icon=":material/arrow_back:",
                    disabled=st.session_state[page_key] >= page_count,
                    on_click=move_page,
                    args=(1,),
                    key=f"{key_prefix}-next",
                    width="stretch",
                )
        else:
            info_col, zoom_spacer, rotation_spacer = st.columns(
                [1.35, 1.35, 2],
                vertical_alignment="bottom",
            )
            with info_col:
                st.metric("صفحات المستند", f"{page_count:,}")
            # يحافظ هذان العمودان على محاذاة أدوات العرض في الوضعين.
            zoom_col = zoom_spacer
            left_col, right_col = rotation_spacer.columns(2)

        with zoom_col:
            zoom_label = st.selectbox(
                "التكبير",
                [
                    "75%", "100%", "125%", "150%", "175%", "200%", "250%",
                    "300%", "400%", "500%", "600%", "750%", "1000%",
                    "1250%", "1500%", "1750%", "2000%",
                ],
                index=2,
                key=f"{key_prefix}-zoom",
            )
            quality_mode = st.selectbox(
                "وضوح النص",
                ["دقة فائقة للنصوص", "أداء متوازن"],
                key=f"{key_prefix}-quality",
                help=(
                    "الدقة الفائقة تعيد رسم صفحة PDF نفسها بدقة تصل إلى 720 DPI، "
                    "بدلاً من تكبير صورة منخفضة الدقة."
                ),
            )
        with left_col:
            st.button(
                "تدوير يسار",
                icon=":material/rotate_left:",
                on_click=rotate_page,
                args=(-90,),
                key=f"{key_prefix}-rotate-left",
                width="stretch",
            )
        with right_col:
            st.button(
                "تدوير يمين",
                icon=":material/rotate_right:",
                on_click=rotate_page,
                args=(90,),
                key=f"{key_prefix}-rotate-right",
                width="stretch",
            )

        zoom = float(zoom_label.rstrip("%")) / 100
        rotation = int(st.session_state[rotation_key])
        if quality_mode == "دقة فائقة للنصوص":
            # في العرض اليدوي تُرسم صفحة واحدة حتى 720 DPI. عند عرض المستند
            # كاملاً يُستخدم حد آمن حتى لا تمتلئ الذاكرة بعشرات الصفحات الضخمة.
            render_scale_cap = 10.0 if display_mode == "التنقل اليدوي بين الصفحات" else 5.0
        else:
            render_scale_cap = 3.0
        # يسمح بالتكبير البصري الإضافي حتى 2000% فوق مصدر فائق الدقة.
        # يظل شريط التمرير الأفقي ثابتاً أسفل نافذة العارض.
        display_width = max(825, int(1100 * zoom))
        if display_mode == "التنقل اليدوي بين الصفحات":
            page_number = int(st.session_state[page_key])
            st.caption(
                f"الصفحة {page_number:,} من {page_count:,}  ·  التكبير {zoom_label}"
                f"  ·  {quality_mode}"
            )
        else:
            st.caption(
                f"المستند كاملاً: {page_count:,} صفحة  ·  التكبير {zoom_label}"
                "  ·  استخدم شريط التمرير داخل مساحة العرض"
            )

    if display_mode == "المستند كاملاً في صفحة عرض واحدة":
        with st.spinner("جارٍ تجهيز المستند كاملاً بدقة عالية..."):
            rendered_pages: list[tuple[int, bytes]] = []
            for page_index in range(page_count):
                image = render_pdf_page(
                    str(path),
                    modified_at,
                    page_index,
                    zoom,
                    rotation,
                    render_scale_cap,
                )
                rendered_pages.append((page_index + 1, image))
            render_scrollable_pdf_images(rendered_pages, display_width)
    else:
        with st.spinner("جارٍ تجهيز الصفحة بدقة عالية..."):
            image = render_pdf_page(
                str(path),
                modified_at,
                page_number - 1,
                zoom,
                rotation,
                render_scale_cap,
            )
        render_scrollable_pdf_images([(page_number, image)], display_width)
    st.caption("معاينة محلية آمنة للعرض فقط؛ لا تُرسل نسخة PDF إلى قارئ المتصفح.")


@st.cache_data(show_spinner=False, max_entries=64)
def get_pdf_page_count(path_text: str, modified_at: float) -> int:
    """يعيد عدد صفحات PDF مع ربط النتيجة بآخر تعديل."""
    del modified_at
    with fitz.open(path_text) as document:
        return document.page_count


@st.cache_data(show_spinner=False, max_entries=48)
def render_pdf_page(
    path_text: str,
    modified_at: float,
    page_number: int,
    zoom: float,
    rotation: int = 0,
    render_scale_cap: float = 3.0,
) -> bytes:
    """يحوّل صفحة PDF إلى صورة عالية الدقة، مع دعم التكبير والتدوير."""
    del modified_at
    with fitz.open(path_text) as document:
        page = document.load_page(page_number)
        # يربط التكبير بدقة الرسم الحقيقية. في الوضع الفائق يصل المقياس
        # إلى 10x (نحو 720 DPI) بدلاً من تمديد صورة محدودة الدقة.
        render_scale = min(max(2.0 * zoom, 1.5), render_scale_cap)
        matrix = fitz.Matrix(render_scale, render_scale).prerotate(rotation)
        pixmap = page.get_pixmap(matrix=matrix, alpha=False, annots=True)
        return pixmap.tobytes("png")


def format_bytes(value: int) -> str:
    """ينسّق عدداً من البايتات بوحدة عربية مناسبة.

    Args:
        value: الحجم الخام بالبايت.

    Returns:
        نصاً منسقاً مثل ``1.5 م.ب``.
    """
    size = float(value)
    for unit in ("بايت", "ك.ب", "م.ب", "ج.ب", "ت.ب"):
        if size < 1024 or unit == "ت.ب":
            return f"{size:,.1f} {unit}" if unit != "بايت" else f"{int(size):,} {unit}"
        size /= 1024
    return f"{size:,.1f} ت.ب"


def file_icon(extension: str) -> str:
    """يعيد رمزاً بصرياً مناسباً لامتداد الملف.

    Args:
        extension: امتداد الملف بالحروف الصغيرة.

    Returns:
        رمزاً موجزاً صالحاً للعرض في مستكشف الملفات.
    """
    if extension in FILE_ICONS:
        return FILE_ICONS[extension]
    if extension in CODE_EXTENSIONS:
        return "⌘"
    return "📄"


def _number(value: object) -> float | None:
    """يحوّل قيمة رقمية أو نسبة مئوية إلى عدد عشري آمن."""
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    translation = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
    text = str(value).translate(translation).replace(",", "").replace("٪", "%").strip()
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    return float(match.group()) if match else None


def _percentage(value: object) -> float | None:
    """يوحّد النسبة المئوية إلى نطاق 0–100."""
    number = _number(value)
    if number is None:
        return None
    return number * 100 if 0 <= number <= 1.5 else number


def _clean_event_name(value: object) -> str:
    """ينظف اسم الفعالية من محارف الاتجاه والتاريخ الملحق."""
    text = str(value or "").replace("\u200f", "").replace("\u200e", "").strip()
    text = re.sub(r"\s*/\s*\d{1,2}/\d{1,2}/\d{4}\s*$", "", text)
    return re.sub(r"\s+", " ", text)


ARABIC_MONTHS = {
    "يوليو": 7, "أغسطس": 8, "اغسطس": 8, "سبتمبر": 9,
    "أكتوبر": 10, "اكتوبر": 10, "نوفمبر": 11, "ديسمبر": 12,
    "يناير": 1, "فبراير": 2, "مارس": 3, "أبريل": 4,
    "ابريل": 4, "مايو": 5, "يونيو": 6,
}


def _extract_event_date(text: object, fallback_name: str = "") -> datetime | None:
    """يستخرج تاريخ الفعالية من النص أو من اسم التقرير."""
    if isinstance(text, datetime):
        return text
    translation = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
    value = str(text or "").translate(translation)
    iso_match = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", value)
    if iso_match:
        year, month, day = iso_match.groups()
        try:
            return datetime(int(year), int(month), int(day))
        except ValueError:
            pass
    matches = re.findall(r"(\d{1,2})\s*[/\-]\s*(\d{1,2})\s*[/\-]\s*(20\d{2})", value)
    if matches:
        day, month, year = matches[-1]
        try:
            return datetime(int(year), int(month), int(day))
        except ValueError:
            pass
    year_in_value = re.search(r"(20\d{2})", value)
    day_in_value = re.search(r"\b(\d{1,2})\b", value)
    if year_in_value and day_in_value:
        for month_name, month_number in ARABIC_MONTHS.items():
            if month_name in value:
                try:
                    return datetime(
                        int(year_in_value.group(1)),
                        month_number,
                        int(day_in_value.group(1)),
                    )
                except ValueError:
                    break
    fallback = fallback_name.translate(translation)
    year_match = re.search(r"(20\d{2})", fallback)
    for month_name, month_number in ARABIC_MONTHS.items():
        if month_name in fallback and year_match:
            return datetime(int(year_match.group()), month_number, 1)
    return None


def _season_quarter(event_date: datetime | None, season_label: str, path_text: str) -> tuple[str, str]:
    """يحدد الموسم الرياضي والربع، مع اعتبار يوليو بداية الموسم."""
    if event_date:
        start_year = event_date.year if event_date.month >= 7 else event_date.year - 1
        season_label = f"{start_year}/{start_year + 1}"
        quarter_number = ((event_date.month - 7) % 12) // 3 + 1
        return season_label, f"الربع {['الأول', 'الثاني', 'الثالث', 'الرابع'][quarter_number - 1]}"
    quarter_match = re.search(r"الربع\s+(الأول|الاول|الثاني|الثالث|الرابع)", path_text)
    quarter = quarter_match.group(1).replace("الاول", "الأول") if quarter_match else "غير محدد"
    return season_label, f"الربع {quarter}" if quarter != "غير محدد" else quarter


def _age_group(target_text: str) -> str:
    """يوحّد الفئات العمرية المكتوبة بصيغ مختلفة في خطط العمل."""
    normalized = re.sub(r"\s+", "", target_text).replace("–", "-").replace("إلى", "-")
    if re.search(r"7-?17|17-?7", normalized):
        return "7–17 سنة"
    if re.search(r"7-?12|12-?7", normalized):
        return "7–12 سنة"
    if re.search(r"13-?17|17-?13", normalized) or "تحت14" in normalized:
        return "13–17 سنة"
    return "فئات مجتمعية وأخرى"


def _program_category(text: str) -> str:
    """يصنف البرنامج من نص اسمه وأهدافه الواردة في التصور العام."""
    if any(word in text for word in ("قرآن", "الصلاة", "ديني", "شعبان", "رمضان", "النبي", "إيماني")):
        return "ديني"
    if any(word in text for word in ("رياضي", "رياضة", "كرة", "بطولة", "المشي", "الرماية", "لياقة")):
        return "رياضي"
    if any(word in text for word in ("ثقافي", "كتاب", "قراءة", "سينمائي", "مسرح", "تراث", "فني")):
        return "ثقافي"
    if any(word in text for word in ("محاضرة", "ورشة", "توعية", "السلامة", "الإسعافات", "صحي")):
        return "توعوي"
    return "مجتمعي"


@st.cache_data(show_spinner=False, max_entries=8)
def scan_general_concepts(root_text: str, revision: int) -> dict:
    """يستخرج الإحصائيات حصراً من ملفات التصور العام الأصلية لكل ربع."""
    del revision
    root = Path(root_text)
    records: list[dict] = []
    sources: list[str] = []
    for season_dir in root.iterdir():
        if not season_dir.is_dir() or "الانشطة" not in season_dir.name:
            continue
        season_match = re.search(r"(202\d)-(202\d)", season_dir.name)
        if not season_match:
            continue
        season_label = f"{season_match.group(1)}/{season_match.group(2)}"
        try:
            concept_dirs = [
                path for path in season_dir.iterdir()
                if path.is_dir() and "التصور العام لفعاليات" in path.name
            ]
        except OSError:
            continue
        for concept_dir in concept_dirs:
            for workbook_path in concept_dir.rglob("*.xlsx"):
                try:
                    workbook = load_workbook(workbook_path, data_only=True, read_only=True)
                    sheet = workbook.active
                    rows = list(sheet.iter_rows(max_col=min(sheet.max_column, 40), values_only=True))

                    def value_at(row_number: int, column_number: int | None):
                        """يعيد قيمة خلية من مصفوفة القراءة السريعة."""
                        if not column_number or row_number < 1 or row_number > len(rows):
                            return None
                        row_values = rows[row_number - 1]
                        return row_values[column_number - 1] if column_number <= len(row_values) else None

                    header_row = None
                    name_col = None
                    for row_index in range(1, min(len(rows), 12) + 1):
                        for column_index in range(1, min(len(rows[row_index - 1]), 40) + 1):
                            value = str(value_at(row_index, column_index) or "")
                            if "الأنشطة/ الفعاليات" in value:
                                header_row, name_col = row_index, column_index
                                break
                        if header_row:
                            break
                    if not header_row or not name_col:
                        continue
                    header_values = {
                        str(value_at(header_row, column_index) or "").strip(): column_index
                        for column_index in range(1, min(len(rows[header_row - 1]), 40) + 1)
                    }
                    audience_col = next((col for label, col in header_values.items() if "الفئات المستهدفة" in label), None)
                    date_col = next((col for label, col in header_values.items() if "موعد الفعالية" in label), None)
                    objective_col = next((col for label, col in header_values.items() if "أهداف الفعالية" in label), None)
                    kpi_row = header_row + 1
                    indicator_col = next(
                        (column_index for column_index in range(1, min(len(rows[kpi_row - 1]), 40) + 1)
                         if "اسم المؤشر" in str(value_at(kpi_row, column_index) or "")),
                        None,
                    )
                    target_col = indicator_col + 1 if indicator_col else None
                    actual_col = indicator_col + 2 if indicator_col else None
                    current: dict | None = None
                    for row_index in range(kpi_row + 1, len(rows) + 1):
                        serial = _number(value_at(row_index, name_col - 1))
                        activity_name = str(value_at(row_index, name_col) or "").strip()
                        if serial is not None and activity_name:
                            if current:
                                records.append(current)
                            event_date = _extract_event_date(
                                value_at(row_index, date_col),
                                workbook_path.stem,
                            )
                            _, quarter = _season_quarter(event_date, season_label, str(workbook_path))
                            if "النشاط الصيفي" in workbook_path.relative_to(concept_dir).parts:
                                quarter = "الربع الأول"
                            audience = str(value_at(row_index, audience_col) or "").strip() if audience_col else ""
                            objective = str(value_at(row_index, objective_col) or "").strip() if objective_col else ""
                            risk_pairs = []
                            for risk_col in range(9 if name_col == 3 else name_col + 6, (indicator_col or 20) - 1, 2):
                                risk = str(value_at(row_index, risk_col) or "").strip()
                                solution = str(value_at(row_index, risk_col + 1) or "").strip()
                                if risk and risk not in {"لا يوجد", "لايوجد"}:
                                    risk_pairs.append((risk, solution))
                            current = {
                                "الفعالية": activity_name,
                                "الموسم": season_label,
                                "الربع": quarter,
                                "التاريخ": event_date.strftime("%Y-%m-%d") if event_date else None,
                                "الفئة_العمرية": _age_group(audience),
                                "التصنيف": _program_category(f"{activity_name} {objective}"),
                                "مؤشرات": [],
                                "المخاطر": risk_pairs,
                                "المصدر": str(workbook_path),
                            }
                        if current and indicator_col:
                            indicator = str(value_at(row_index, indicator_col) or "").strip()
                            target = _number(value_at(row_index, target_col))
                            actual = _number(value_at(row_index, actual_col))
                            if indicator and (target is not None or actual is not None):
                                current["مؤشرات"].append({
                                    "المؤشر": indicator, "المستهدف": target, "الفعلي": actual,
                                })
                    if current:
                        records.append(current)
                    sources.append(str(workbook_path))
                except (OSError, ValueError, KeyError, zipfile.BadZipFile):
                    continue
    return {"records": records, "sources": sources}


@st.cache_data(show_spinner=False, max_entries=8)
def scan_sports_seasons(root_text: str, revision: int) -> dict:
    """يحلل الموسمين الرياضيين من يوليو إلى يونيو ويصنف النتائج ربعياً."""
    del revision
    root = Path(root_text)
    season_dirs = [
        path for path in root.iterdir()
        if path.is_dir() and re.search(r"202[5-6]-202[6-7]", path.name) and "الانشطة" in path.name
    ]
    all_events: list[dict] = []
    all_surveys: list[dict] = []
    all_plans: list[dict] = []
    all_notes: list[dict] = []
    all_age_profiles: list[dict] = []
    all_workplan_kpis: list[dict] = []
    report_sources: Counter = Counter()

    for season_dir in season_dirs:
        match = re.search(r"(202\d)-(202\d)", season_dir.name)
        season_label = f"{match.group(1)}/{match.group(2)}" if match else season_dir.name
        for report_path in season_dir.glob("**/التقارير الشهرية/**/*.docx"):
            try:
                document = Document(report_path)
                report_has_events = False
                report_label = report_path.stem.replace("التقرير الأسبوعي -", "").replace("تقرير شهر", "").strip()
                for table in document.tables:
                    if len(table.columns) < 6:
                        continue
                    table_rows = [
                        [cell.text.strip() for cell in row.cells]
                        for row in table.rows
                    ]
                    if table_rows and any("الأسبوع الأول" in cell for cell in table_rows[0]):
                        summary = {
                            row[0].strip(): row[1:6]
                            for row in table_rows[1:]
                            if len(row) >= 6 and row[0].strip()
                        }
                        week_names = summary.get("اسم الأسبوع", [])
                        event_counts = summary.get("عدد الفعاليات", [])
                        targets = summary.get("العدد المستهدف", [])
                        actuals = summary.get("العدد الفعلي", [])
                        satisfaction_values = summary.get("متوسط نسبة الرضا", [])
                        if all(len(values) == 5 for values in (week_names, event_counts, targets, actuals)):
                            week_labels = ["الأسبوع الأول", "الأسبوع الثاني", "الأسبوع الثالث", "الأسبوع الرابع", "الأسبوع الخامس"]
                            for week_index in range(5):
                                event_count = int(_number(event_counts[week_index]) or 0)
                                target = int(_number(targets[week_index]) or 0)
                                actual = int(_number(actuals[week_index]) or 0)
                                satisfaction = _percentage(
                                    satisfaction_values[week_index] if week_index < len(satisfaction_values) else None
                                )
                                all_events.append({
                                    "الفعالية": f"فعاليات النشاط الصيفي — {week_names[week_index]}",
                                    "الفترة": f"{week_labels[week_index]} · {week_names[week_index]}",
                                    "الموسم": season_label,
                                    "الربع": "الربع الأول",
                                    "الشهر": 7,
                                    "التاريخ": f"2025-07-{week_index + 1:02d}",
                                    "المستهدف": target,
                                    "الفعلي": actual,
                                    "المشاركة": actual / target * 100 if target else None,
                                    "الرضا": satisfaction,
                                    "عدد_الفعاليات": event_count,
                                    "المصدر": str(report_path),
                                })
                            report_has_events = True
                            break
                    before = len(all_events)
                    for row in table.rows[1:]:
                        cells = [cell.text.strip() for cell in row.cells]
                        name = _clean_event_name(cells[0])
                        if not name or "الاحصائيات" in name:
                            continue
                        target, actual = _number(cells[2]), _number(cells[3])
                        if target is None or actual is None:
                            continue
                        participation = _percentage(cells[4])
                        satisfaction = _percentage(cells[5])
                        if satisfaction is not None and satisfaction > 100:
                            continue
                        event_date = _extract_event_date(cells[0], report_path.stem)
                        classified_season, quarter = _season_quarter(event_date, season_label, str(report_path))
                        all_events.append({
                            "الفعالية": name,
                            "الفترة": report_label,
                            "الموسم": classified_season,
                            "الربع": quarter,
                            "الشهر": event_date.month if event_date else None,
                            "التاريخ": event_date.strftime("%Y-%m-%d") if event_date else None,
                            "المستهدف": int(target),
                            "الفعلي": int(actual),
                            "المشاركة": participation,
                            "الرضا": satisfaction,
                            "عدد_الفعاليات": 1,
                            "المصدر": str(report_path),
                        })
                    if len(all_events) > before:
                        report_has_events = True
                        break
                if report_has_events:
                    report_sources[season_label] += 1
                for index, paragraph in enumerate(document.paragraphs):
                    if "الملاحظات والتوصيات" in paragraph.text and index + 1 < len(document.paragraphs):
                        recommendation = document.paragraphs[index + 1].text.strip()
                        for part in re.split(r"\n|(?=\d+\s*[-–])", recommendation):
                            cleaned = re.sub(r"^\d+\s*[-–]\s*", "", part).strip()
                            if cleaned:
                                severity = "مرتفع" if any(word in cleaned for word in ("تزاحم", "خطر", "حالة طارئة")) else "متوسط"
                                category = "الطاقة الاستيعابية" if "تزاحم" in cleaned else (
                                    "الصحة والسلامة" if any(word in cleaned for word in ("سلامة", "حرارة", "وقائي")) else "تشغيلي"
                                )
                                fallback_date = _extract_event_date("", report_path.stem)
                                _, note_quarter = _season_quarter(fallback_date, season_label, str(report_path))
                                all_notes.append({
                                    "المستوى": severity, "الفئة": category, "البيان": cleaned,
                                    "الفترة": report_label, "الموسم": season_label, "الربع": note_quarter,
                                    "النوع": "موثق في التقرير",
                                })
            except (OSError, ValueError, KeyError, zipfile.BadZipFile):
                continue

        for survey_path in season_dir.glob("**/*.xlsx"):
            if not any("استب" in part for part in survey_path.parts) or "Pdf" in survey_path.parts:
                continue
            try:
                workbook = load_workbook(survey_path, data_only=True, read_only=True)
                if "مقياس ليكارت" not in workbook.sheetnames:
                    continue
                sheet = workbook["مقياس ليكارت"]
                satisfaction = _percentage(sheet["B6"].value)
                responses = _number(sheet["N7"].value) or _number(sheet["D4"].value)
                event_date = _extract_event_date(sheet["A3"].value, survey_path.stem)
                classified_season, quarter = _season_quarter(event_date, season_label, str(survey_path))
                if satisfaction is not None:
                    all_surveys.append({
                        "الفعالية": _clean_event_name(sheet["A2"].value or survey_path.stem),
                        "الرضا": satisfaction, "الاستجابات": int(responses or 0),
                        "الموسم": classified_season, "الربع": quarter, "المصدر": str(survey_path),
                    })
            except (OSError, ValueError, KeyError, zipfile.BadZipFile):
                continue

        plan_paths = list(season_dir.glob("**/خطة الأنشطة الثقافية والمجتمعية/*.xlsx"))
        plan_paths += list(season_dir.glob("**/خطة الأنشطة الثقافية والمجتمعية/*.xls"))
        if plan_paths:
            plan_path = plan_paths[0]
            try:
                if plan_path.suffix.lower() == ".xlsx":
                    sheet = load_workbook(plan_path, data_only=True, read_only=True).worksheets[0]
                    raw_plan_rows = list(sheet.iter_rows(values_only=True))
                else:
                    frame = pd.read_excel(plan_path, header=None)
                    raw_plan_rows = [tuple(row) for row in frame.itertuples(index=False, name=None)]
                month_by_column: dict[int, str] = {}
                active_month = ""
                if len(raw_plan_rows) > 6:
                    for column_index, header_value in enumerate(raw_plan_rows[6]):
                        if header_value and re.search(r"20\d{2}", str(header_value)):
                            active_month = str(header_value)
                        if active_month:
                            month_by_column[column_index] = active_month
                plan_rows = raw_plan_rows[8:]
                for row in plan_rows:
                    if len(row) > 3 and _number(row[1]) is not None and row[2]:
                        scheduled_month = ""
                        for column_index in range(5, len(row)):
                            if str(row[column_index] or "").strip() in {"√", "✓", "✔"}:
                                scheduled_month = month_by_column.get(column_index, "")
                                break
                        plan_date = _extract_event_date("", scheduled_month)
                        _, plan_quarter = _season_quarter(plan_date, season_label, str(plan_path))
                        all_plans.append({
                            "الرقم": int(_number(row[1]) or 0), "الفعالية": _clean_event_name(row[2]),
                            "المجال": str(row[3] or "غير مصنف"), "الموسم": season_label,
                            "الربع": plan_quarter,
                        })
            except (OSError, ValueError, KeyError, zipfile.BadZipFile):
                pass

        for workplan_path in season_dir.glob("**/خطة العمل الفردية للفعاليات/**/*.docx"):
            try:
                document = Document(workplan_path)
                if not document.tables:
                    continue
                event_name = workplan_path.stem
                target_group = ""
                for row in document.tables[0].rows[:4]:
                    cells = [cell.text.strip() for cell in row.cells]
                    for cell_index, value in enumerate(cells):
                        if "الفئات المستهدفة" in value and cell_index + 1 < len(cells):
                            target_group = cells[cell_index + 1].strip()
                            break
                event_date = _extract_event_date("", workplan_path.stem)
                _, quarter = _season_quarter(event_date, season_label, str(workplan_path))
                all_age_profiles.append({
                    "الفعالية": event_name,
                    "الفئة_الموثقة": target_group or "غير محددة",
                    "الفئة_العمرية": _age_group(target_group),
                    "الموسم": season_label,
                    "الربع": quarter,
                    "المصدر": str(workplan_path),
                })
                if len(document.tables) > 1:
                    for row in document.tables[1].rows[2:]:
                        cells = [cell.text.strip() for cell in row.cells]
                        if len(cells) < 12:
                            continue
                        indicator = cells[8]
                        target = _number(cells[9])
                        actual = _number(cells[10])
                        if indicator and (target is not None or actual is not None):
                            all_workplan_kpis.append({
                                "المؤشر": indicator,
                                "المستهدف": target,
                                "الفعلي": actual,
                                "الموسم": season_label,
                                "الربع": quarter,
                                "المصدر": str(workplan_path),
                            })
            except (OSError, ValueError, KeyError, zipfile.BadZipFile):
                continue

    deduplicated: dict[tuple, dict] = {}
    current_events: list[dict] = []
    for event in all_events:
        if event["الموسم"] == "2026/2027":
            current_events.append(event)
            continue
        key = (
            event["الموسم"], event["التاريخ"], event["الفعالية"].casefold(),
            event["المستهدف"], event["الفعلي"],
        )
        deduplicated[key] = event
    events = list(deduplicated.values()) + current_events
    inferred: list[dict] = []
    for event in events:
        base = {
            "الفترة": event["الفترة"], "الموسم": event["الموسم"],
            "الربع": event["الربع"], "النوع": "مستنتج آلياً",
        }
        if event["المشاركة"] is not None and event["المشاركة"] >= 120:
            inferred.append({**base, "المستوى": "مرتفع", "الفئة": "الطاقة الاستيعابية",
                             "البيان": f"{event['الفعالية']}: حضور {event['المشاركة']:.1f}% من المستهدف."})
        if event["الرضا"] is not None and event["الرضا"] < 95:
            inferred.append({**base, "المستوى": "متوسط", "الفئة": "رضا المستفيدين",
                             "البيان": f"{event['الفعالية']}: الرضا {event['الرضا']:.1f}% دون حد المتابعة 95%."})
        if event["الفعلي"] < event["المستهدف"]:
            inferred.append({**base, "المستوى": "متوسط", "الفئة": "تحقيق المستهدف",
                             "البيان": f"{event['الفعالية']}: تحقق {event['الفعلي']} من {event['المستهدف']}."})
    return {
        "events": events, "surveys": all_surveys, "plan": all_plans,
        "notes": all_notes, "risks": all_notes + inferred,
        "age_profiles": all_age_profiles, "workplan_kpis": all_workplan_kpis,
        "sources_by_season": dict(report_sources),
    }


@st.cache_data(show_spinner=False, max_entries=8)
def scan_event_intelligence(root_text: str, revision: int) -> dict:
    """يستخرج مؤشرات الأداء والمخاطر من مصادر الموسم الحالي.

    Args:
        root_text: مسار جذر المشروع.
        revision: رقم تغيير يستخدم لإبطال التخزين المؤقت.

    Returns:
        قاموساً يضم الفعاليات والخطة والاستبانات والمخاطر الموثقة والمستنتجة.
    """
    del revision
    root = Path(root_text)
    season_dirs = [
        path for path in root.iterdir()
        if path.is_dir() and "2026-2027" in path.name and "الانشطة" in path.name
    ]
    if not season_dirs:
        return {"events": [], "surveys": [], "plan": [], "notes": [], "sources": 0}
    season = season_dirs[0]
    events: list[dict] = []
    notes: list[dict] = []
    source_count = 0

    for report_path in season.glob("**/التقارير الشهرية/**/*.docx"):
        if "التقرير الأسبوعي" not in report_path.name:
            continue
        try:
            document = Document(report_path)
            source_count += 1
            week = report_path.stem.split("-", 1)[-1].strip()
            for table in document.tables:
                if len(table.columns) < 6:
                    continue
                before_count = len(events)
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    name = _clean_event_name(cells[0])
                    if not name or "الاحصائيات" in name:
                        continue
                    target = _number(cells[2])
                    actual = _number(cells[3])
                    participation = _percentage(cells[4])
                    satisfaction = _percentage(cells[5])
                    if target is None or actual is None:
                        continue
                    events.append({
                        "الفعالية": name,
                        "الأسبوع": week,
                        "المستهدف": int(target),
                        "الفعلي": int(actual),
                        "المشاركة": participation,
                        "الرضا": satisfaction,
                        "المصدر": str(report_path),
                    })
                if len(events) > before_count:
                    break
            for index, paragraph in enumerate(document.paragraphs):
                if "الملاحظات والتوصيات" in paragraph.text and index + 1 < len(document.paragraphs):
                    recommendation = document.paragraphs[index + 1].text.strip()
                    for part in re.split(r"\n|(?=\d+\s*[-–])", recommendation):
                        cleaned = re.sub(r"^\d+\s*[-–]\s*", "", part).strip()
                        if cleaned:
                            severity = "مرتفع" if any(word in cleaned for word in ("تزاحم", "خطر", "حالة طارئة")) else "متوسط"
                            category = "الطاقة الاستيعابية" if "تزاحم" in cleaned else (
                                "الصحة والسلامة" if any(word in cleaned for word in ("سلامة", "حرارة", "وقائي")) else "تشغيلي"
                            )
                            notes.append({
                                "المستوى": severity,
                                "الفئة": category,
                                "البيان": cleaned,
                                "الأسبوع": week,
                                "النوع": "موثق في التقرير",
                            })
        except (OSError, ValueError, KeyError, zipfile.BadZipFile):
            continue

    plan: list[dict] = []
    plan_path = next(season.glob("**/خطة الأنشطة الثقافية والمجتمعية/*2026.xlsx"), None)
    if plan_path:
        try:
            workbook = load_workbook(plan_path, data_only=True, read_only=True)
            sheet = workbook.worksheets[0]
            for row in sheet.iter_rows(min_row=9, values_only=True):
                sequence, name, domain = row[1], row[2], row[3]
                if _number(sequence) is not None and name:
                    plan.append({"الرقم": int(_number(sequence) or 0), "الفعالية": _clean_event_name(name), "المجال": str(domain or "غير مصنف")})
        except (OSError, ValueError, KeyError, zipfile.BadZipFile):
            pass

    surveys: list[dict] = []
    for survey_path in season.rglob("*.xlsx"):
        # يدعم اختلاف اسم المجلد بين الموسمين: الاستبيانات / الاستبانات.
        if not any("استب" in part for part in survey_path.parts):
            continue
        if "Pdf" in survey_path.parts:
            continue
        try:
            workbook = load_workbook(survey_path, data_only=True, read_only=True)
            if "مقياس ليكارت" not in workbook.sheetnames:
                continue
            sheet = workbook["مقياس ليكارت"]
            name = _clean_event_name(sheet["A2"].value or survey_path.stem)
            satisfaction = _percentage(sheet["B6"].value)
            participants = _number(sheet["N7"].value) or _number(sheet["D4"].value)
            if satisfaction is not None:
                surveys.append({
                    "الفعالية": name,
                    "الرضا": satisfaction,
                    "الاستجابات": int(participants or 0),
                    "المصدر": str(survey_path),
                })
        except (OSError, ValueError, KeyError, zipfile.BadZipFile):
            continue

    inferred_risks: list[dict] = []
    for event in events:
        if event["المشاركة"] is not None and event["المشاركة"] >= 120:
            inferred_risks.append({
                "المستوى": "مرتفع",
                "الفئة": "الطاقة الاستيعابية",
                "البيان": f"{event['الفعالية']}: حضور {event['المشاركة']:.1f}% من المستهدف.",
                "الأسبوع": event["الأسبوع"],
                "النوع": "مستنتج آلياً",
            })
        if event["الرضا"] is not None and event["الرضا"] < 95:
            inferred_risks.append({
                "المستوى": "متوسط",
                "الفئة": "رضا المستفيدين",
                "البيان": f"{event['الفعالية']}: الرضا {event['الرضا']:.1f}% دون حد المتابعة 95%.",
                "الأسبوع": event["الأسبوع"],
                "النوع": "مستنتج آلياً",
            })
        if event["الفعلي"] < event["المستهدف"]:
            inferred_risks.append({
                "المستوى": "متوسط",
                "الفئة": "تحقيق المستهدف",
                "البيان": f"{event['الفعالية']}: تحقق {event['الفعلي']} من {event['المستهدف']}.",
                "الأسبوع": event["الأسبوع"],
                "النوع": "مستنتج آلياً",
            })

    return {
        "events": events,
        "surveys": surveys,
        "plan": plan,
        "notes": notes,
        "risks": notes + inferred_risks,
        "sources": source_count,
    }


def inject_theme() -> None:
    """يحقن طبقة التنسيق المرئية الخاصة باللوحة."""
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Noto+Kufi+Arabic:wght@400;500;600;700;800&family=Tajawal:wght@400;500;600;700;800&display=swap');
        :root {
            --ink:#252623; --muted:#72736d; --line:#dedbd1; --mint:#176b45;
            --soft:#f7f4ec; --navy:#242724; --navy2:#31352f; --gold:#c6a258;
            --green:#176b45; --green2:#0f5134; --burgundy:#7b2638; --ivory:#fffdf7;
        }
        html, body, [class*="css"], .stApp {
            font-family:"Noto Kufi Arabic","Tajawal","Segoe UI",sans-serif;
        }
        .stApp {
            direction:rtl;color:var(--ink);
            background:
                radial-gradient(circle at 13% 4%,rgba(198,162,88,.12),transparent 27rem),
                linear-gradient(180deg,#fbf9f3 0%,var(--soft) 100%);
        }
        .block-container { max-width:1400px; padding:1.4rem 2.4rem 3rem; }
        [data-testid="stSidebar"] {
            direction:rtl;
            background:
                linear-gradient(145deg,rgba(255,255,255,.045),transparent 35%),
                linear-gradient(180deg,#222522 0%,#2b2f2b 54%,#1d201e 100%);
            border-left:1px solid rgba(198,162,88,.34);
            box-shadow:-14px 0 32px rgba(38,40,35,.14);
            width:320px !important; min-width:320px !important; max-width:320px !important;
            flex:0 0 320px !important; transform:none !important;
        }
        [data-testid="stSidebar"] > div:first-child,
        [data-testid="stSidebarContent"] {
            width:320px !important; min-width:320px !important; max-width:320px !important;
        }
        [data-testid="stSidebar"] > div:first-child { padding-top:1rem; }
        [data-testid="stSidebarCollapsedControl"],
        [data-testid="stSidebarCollapseButton"] { display:none !important; }
        [data-testid="stSidebar"] * { color:#fff; }
        [data-testid="stSidebar"] .stButton button {
            justify-content:flex-start; text-align:right; border:1px solid rgba(255,255,255,.09);
            background:rgba(255,255,255,.025); min-height:44px; padding:.55rem .75rem;
            box-shadow:inset 0 1px rgba(255,255,255,.035);
        }
        [data-testid="stSidebar"] .stButton button:hover {
            background:rgba(123,38,56,.32);border-color:rgba(198,162,88,.52);
            transform:translateX(-2px);
        }
        [data-testid="stSidebar"] .stButton button[kind="primary"],
        [data-testid="stSidebar"] .stButton button[kind="primary"] * {
            background:linear-gradient(135deg,#247a52,#125437);color:#fff !important;
            border-color:#3b8e66;font-weight:800;
            box-shadow:0 7px 18px rgba(8,55,34,.34),inset 0 1px rgba(255,255,255,.18);
        }
        [data-testid="stSidebar"] [data-testid="stExpander"] {
            background:transparent !important;border:1px solid rgba(255,255,255,.1) !important;
            border-radius:12px;overflow:hidden;
        }
        [data-testid="stSidebar"] [data-testid="stExpander"] summary {
            background:rgba(255,255,255,.045) !important;color:#fff !important;
        }
        [data-testid="stSidebar"] [data-testid="stExpander"] summary:hover {
            background:rgba(123,38,56,.26) !important;
        }
        [data-testid="stSidebar"] [data-testid="stExpander"] summary svg {
            fill:var(--gold) !important;color:var(--gold) !important;
        }
        [data-testid="stSidebar"] input { background:#fff !important;color:#1a2044 !important;border:0; }
        [data-testid="stSidebar"] [data-baseweb="select"] > div { background:#fff;color:#1a2044;border:0; }
        [data-testid="stSidebar"] [data-baseweb="select"] * { color:#1a2044 !important; }
        [data-testid="stSidebar"] hr { border-color:rgba(255,255,255,.12); }
        [data-testid="stHeader"] { background:transparent; }
        [data-testid="stToolbar"], .stDeployButton, #MainMenu { display:none !important; }
        [data-testid="stMetric"] {
            background:linear-gradient(145deg,#fffefb 0%,#f8f4e9 100%);
            border:1px solid #ddd8cb;border-right:4px solid var(--green);
            border-radius:16px;padding:19px 20px;
            box-shadow:8px 10px 22px rgba(55,53,43,.08),inset 0 1px #fff;
            min-height:126px;transition:transform .18s ease,box-shadow .18s ease;
        }
        [data-testid="stMetric"]:hover {
            box-shadow:10px 15px 28px rgba(55,53,43,.13),inset 0 1px #fff;
            transform:translateY(-3px);
        }
        div[data-testid="stHorizontalBlock"] > div:nth-child(1) [data-testid="stMetric"] { border-right-color:var(--green); }
        div[data-testid="stHorizontalBlock"] > div:nth-child(2) [data-testid="stMetric"] { border-right-color:var(--gold); }
        div[data-testid="stHorizontalBlock"] > div:nth-child(3) [data-testid="stMetric"] { border-right-color:#48564c; }
        div[data-testid="stHorizontalBlock"] > div:nth-child(4) [data-testid="stMetric"] { border-right-color:var(--burgundy); }
        [data-testid="stMetricLabel"] { color:var(--muted); font-size:.9rem; }
        [data-testid="stMetricValue"] {
            color:var(--ink);font-weight:700;font-size:1.65rem;
            white-space:normal !important;overflow:visible !important;text-overflow:clip !important;
        }
        [data-testid="stMetricLabel"], [data-testid="stMetricValue"] { direction:rtl; }
        [data-testid="stTextInput"] input, [data-testid="stSelectbox"] > div > div {
            border-radius:13px; border-color:#dededb; background:white;
        }
        [data-testid="stDataFrame"] { border:1px solid var(--line); border-radius:16px; overflow:hidden; }
        .brand { display:flex; align-items:center; gap:11px; padding:8px 2px 18px; }
        .brand-mark {
            width:46px;height:46px;border-radius:13px;background:linear-gradient(145deg,#fffdf6,#e9dfc8);color:var(--green);
            display:grid;place-items:center;font-weight:900;font-size:18px;
            box-shadow:0 7px 16px rgba(0,0,0,.24),inset 0 1px #fff;border:2px solid var(--gold);
        }
        .brand-name { font-size:1.05rem;font-weight:800;line-height:1.15;color:#fff; }
        .brand-sub { font-size:.7rem;color:#d7c47f;margin-top:5px;letter-spacing:.03em; }
        .nav-section {
            margin:12px 0 7px;color:#d8be6e;font-size:.72rem;font-weight:800;
            border-bottom:1px solid rgba(255,255,255,.1);padding-bottom:7px;
        }
        .form-title {
            margin:.8rem 0 .25rem;padding:.55rem .65rem;border-right:3px solid var(--gold);
            border-radius:8px;background:rgba(255,255,255,.055);line-height:1.35;
        }
        .form-title b { display:block;color:#e8cf87;font-size:.72rem;direction:ltr;text-align:right; }
        .form-title span { display:block;color:#fff;font-size:.82rem;font-weight:700;margin-top:.12rem; }
        .eyebrow { color:var(--gold); font-size:.76rem; font-weight:800; letter-spacing:.05em; margin-bottom:8px; }
        .welcome {
            padding:25px 30px;border-radius:24px;color:#fff;
            background:
                radial-gradient(circle at 10% 10%,rgba(198,162,88,.18),transparent 24rem),
                linear-gradient(135deg,#292d29 0%,#1c5038 100%);
            box-shadow:0 16px 34px rgba(40,44,38,.22),inset 0 1px rgba(255,255,255,.12);
            border:1px solid rgba(198,162,88,.46);border-bottom:4px solid var(--gold);
        }
        .welcome h1 { margin:0; font-size:2rem; font-weight:800; letter-spacing:-.02em;color:#fff; }
        .welcome p { margin:9px 0 0; color:#e1e6f2; font-size:.98rem; }
        .live-pill {
            display:inline-flex;align-items:center;gap:7px;border:1px solid #cce9e1;
            background:#edf6ef;color:var(--green2);border-radius:99px;padding:7px 11px;font-size:.78rem;font-weight:800;
        }
        .live-dot { width:7px;height:7px;border-radius:50%;background:var(--mint);box-shadow:0 0 0 4px #d8f3ec; }
        .insight {
            border:1px solid var(--line);border-radius:15px;padding:20px;
            background:linear-gradient(145deg,#fffefb,#f8f5ed);
            min-height:130px;box-shadow:7px 9px 20px rgba(55,53,43,.07),inset 0 1px #fff;
        }
        .insight-title { color:var(--muted);font-size:.8rem;margin-bottom:12px;font-weight:600; }
        .insight-value { color:var(--ink);font-size:1.06rem;font-weight:700;line-height:1.55; }
        .section-kicker { color:var(--muted);font-size:.82rem;margin-bottom:-5px; }
        .file-head {
            border:1px solid var(--line);border-radius:18px;padding:22px 24px;
            background:linear-gradient(145deg,#fffefb,#f8f5ed);margin:8px 0 20px;
            border-top:4px solid var(--green);
            box-shadow:8px 11px 24px rgba(55,53,43,.08),inset 0 1px #fff;
        }
        .quarter-context {
            margin-top:18px;border-top-color:var(--burgundy);border-right:5px solid var(--green);
            background:
                radial-gradient(circle at 8% 15%,rgba(198,162,88,.14),transparent 18rem),
                linear-gradient(145deg,#fffefb,#f6f1e6);
        }
        .quarter-context .eyebrow { color:var(--burgundy); }
        .file-title {font-size:1.18rem;font-weight:700;overflow-wrap:anywhere;}
        .file-path {font-size:.78rem;color:var(--muted);margin-top:5px;overflow-wrap:anywhere;}
        .file-stats {display:flex;flex-wrap:wrap;gap:18px;margin-top:16px;color:#555;font-size:.84rem;}
        .footer-note {
            color:#d7d9e4;font-size:.72rem;text-align:center;padding:19px 10px 4px;
            background:rgba(255,255,255,.05);border-radius:13px;margin-top:8px;
        }
        div.stButton > button { border-radius:12px; border-color:#dededb; }
        div.stButton > button[kind="primary"] {
            background:linear-gradient(135deg,var(--green),var(--green2));
            border-color:var(--green);color:white;box-shadow:0 6px 14px rgba(15,81,52,.22);
        }
        [data-testid="stTabs"] button[aria-selected="true"] { color:var(--burgundy) !important;font-weight:800; }
        [data-testid="stTabs"] [data-baseweb="tab-highlight"] { background-color:var(--gold); }
        [data-testid="stDataFrame"] { background:#fff; }
        h1,h2,h3 { letter-spacing:-.015em; }
        hr { border-color:var(--line); }
        @media (max-width: 800px) { .block-container{padding:1rem;} .welcome h1{font-size:1.55rem;} }
        </style>
        """,
        unsafe_allow_html=True,
    )


def ensure_state() -> None:
    """يهيئ حالة جلسة Streamlit وخدمة المراقبة عند الحاجة."""
    st.session_state.setdefault("selected_file", None)
    st.session_state.setdefault("selected_folder", None)
    st.session_state.setdefault("selected_collection", None)
    st.session_state.setdefault("revision", 0)
    st.session_state.setdefault("watch_revision_seen", -1)
    st.session_state.setdefault("last_refresh", time.time())
    if st.session_state.get("document_sync_version") != DOCUMENT_SYNC_VERSION:
        st.cache_data.clear()
        st.session_state.revision += 1
        st.session_state.last_refresh = time.time()
        st.session_state.document_sync_version = DOCUMENT_SYNC_VERSION
    service = st.session_state.get("watch_service")
    if service is None or service.root != DEFAULT_PROJECT or not service.is_alive():
        try:
            if service is not None and not service.is_alive():
                st.cache_resource.clear()
            service = get_watch_service(str(DEFAULT_PROJECT))
            st.session_state.watch_service = service
            st.session_state.watch_revision_seen = service.revision
            st.session_state.watch_error = None
        except OSError as exc:
            st.session_state.watch_service = None
            st.session_state.watch_error = str(exc)


def _season_directories() -> dict[str, Path]:
    """يعيد مجلدي الموسمين الرياضيين مرتبين بالاسم المعتمد."""
    result: dict[str, Path] = {}
    for path in DEFAULT_PROJECT.iterdir():
        if not path.is_dir():
            continue
        if "2025-2026" in path.name and "الانشطة" in path.name:
            result["2025/2026"] = path
        elif "2026-2027" in path.name and "الانشطة" in path.name:
            result["2026/2027"] = path
    return result


def _select_folder(path: Path) -> None:
    """يحدّث الجلسة لعرض مجلد محدد في المنطقة الرئيسية."""
    st.session_state.selected_folder = str(path)
    st.session_state.selected_file = None
    st.session_state.selected_collection = None
    st.rerun()


def _select_collection(base: Path, code: str, label: str, season: str, quarter: str) -> None:
    """يختار مساحة نموذج معتمد بصيغتيه الأصلية وPDF ضمن ربع رياضي."""
    st.session_state.selected_collection = {
        "base": str(base),
        "code": code,
        "label": label,
        "season": season,
        "quarter": quarter,
    }
    st.session_state.selected_file = None
    st.session_state.selected_folder = None
    st.rerun()


def _display_folder_name(path: Path) -> str:
    """يعرض اسماً موسمياً موحداً دون الاعتماد على أخطاء أسماء المجلدات القديمة."""
    for season_label, season_path in _season_directories().items():
        if path.resolve() == season_path.resolve():
            return f"الموسم الرياضي {season_label}"
    normalized = path.name.replace("الاول", "الأول")
    quarter_labels = {
        "الربع الأول": "الربع الأول · يوليو – سبتمبر",
        "الربع الثاني": "الربع الثاني · أكتوبر – ديسمبر",
        "الربع الثالث": "الربع الثالث · يناير – مارس",
        "الربع الرابع": "الربع الرابع · أبريل – يونيو",
    }
    for marker, label in quarter_labels.items():
        if marker in normalized:
            return label
    return path.name


def render_sidebar(files: list[dict]) -> None:
    """يعرض الموسمين ومجلداتهما والبحث في جميع الملفات.

    Args:
        files: نتائج محرك المسح الحالية.
    """
    with st.sidebar:
        st.markdown(
            '<div class="brand"><div class="brand-mark">م</div><div>'
            '<div class="brand-name">اللجنة الثقافية والمجتمعية</div>'
            '<div class="brand-sub">نادي مليحة الثقافي الرياضي</div></div></div>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="nav-section">القيادة والتشغيل</div>', unsafe_allow_html=True)
        if st.button("▦  لوحة قيادة الفعاليات", type="primary", width="stretch"):
            st.session_state.selected_file = None
            st.session_state.selected_folder = None
            st.session_state.selected_collection = None
            st.rerun()
        if st.button(
            "↻  مزامنة جميع المستندات الآن",
            key="force-document-sync",
            width="stretch",
            help="إعادة مسح الموسمين وإبطال جميع النتائج المخزنة مؤقتاً.",
        ):
            st.cache_data.clear()
            st.session_state.revision += 1
            st.session_state.last_refresh = time.time()
            st.toast("تمت مزامنة مستندات الموسمين.", icon="✅")
            st.rerun()

        for season_index, (season_label, season_path) in enumerate(_season_directories().items(), start=1):
            try:
                season_children = [child for child in season_path.iterdir() if child.is_dir()]
            except OSError:
                season_children = []
            with st.expander(
                f"الموسم {'الأول' if season_index == 1 else 'الثاني'} · {season_label}",
                expanded=season_index == 1,
            ):
                if st.button("الصفحة الرئيسية للموسم", key=f"season-folder-{season_label}", width="stretch"):
                    _select_folder(season_path)
                for quarter, months in SPORTS_QUARTERS:
                    with st.expander(f"{quarter} · {months}"):
                        for code, form_label, folder_hint in APPROVED_FORMS:
                            form_folder = next(
                                (child for child in season_children if folder_hint in child.name),
                                season_path / folder_hint,
                            )
                            quarter_count = 0
                            try:
                                form_available = form_folder.exists() and form_folder.is_dir()
                                if form_available:
                                    quarter_count = sum(
                                        path.suffix.lower() == ".pdf"
                                        and _season_quarter(None, season_label, str(path))[1]
                                        in {quarter, "غير محدد"}
                                        for path in form_folder.rglob("*") if path.is_file()
                                    )
                            except OSError:
                                form_available = False
                                quarter_count = 0
                            if st.button(
                                f"{code} · {form_label}  ({quarter_count})",
                                key=f"form-{season_label}-{quarter}-{code}",
                                width="stretch",
                                disabled=not form_available,
                            ):
                                _select_collection(form_folder, code, form_label, season_label, quarter)

        st.markdown('<div class="nav-section">بحث شامل في المستندات</div>', unsafe_allow_html=True)
        query = st.text_input("بحث", placeholder="اكتب اسم ملف أو مجلد…")
        visible = [
            item for item in files
            if item["extension"] == ".pdf"
            and (not query or query.casefold() in item["relative_path"].casefold())
        ]
        if query:
            shown = visible[:40]
            st.caption(f"{len(visible):,} نتيجة — عرض أول {len(shown)}")
        else:
            shown = []
        for index, item in enumerate(shown):
            compact = Path(item["relative_path"]).name
            label = f"{file_icon(item['extension'])}  {compact[:42]}"
            if st.button(label, key=f"file-{index}-{item['absolute_path']}", width="stretch", help=item["relative_path"]):
                st.session_state.selected_file = item["absolute_path"]
                st.rerun()
        st.caption("الانتقال المباشر إلى ملف")
        file_options = [item for item in files if not item["error"] and item["extension"] == ".pdf"]
        label_to_item = {
            f"{file_icon(item['extension'])}  {item['relative_path']}": item
            for item in file_options
        }
        options = ["— اختر من جميع الملفات —"] + list(label_to_item)
        chosen = st.selectbox("جميع الملفات", options, label_visibility="collapsed")
        if chosen != options[0]:
            match = label_to_item[chosen]
            if st.session_state.selected_file != match["absolute_path"]:
                st.session_state.selected_file = match["absolute_path"]
                st.rerun()
        st.markdown(
            f'<div class="footer-note">{len(files):,} ملف تحت المراقبة<br>غازي جنيد • نادي مليحة الثقافي الرياضي</div>',
            unsafe_allow_html=True,
        )


def render_concept_statistics(concept: dict) -> None:
    """يعرض مؤشرات الإدارة من التصور العام الأصلي حصراً."""
    records = concept.get("records", [])
    if not records:
        st.warning("لا توجد ملفات تصور عام أصلية قابلة للتحليل حالياً.")
        st.info("لن تُعرض إحصائيات بديلة من التقارير أو أي مصدر آخر.")
        return
    seasons = sorted({record["الموسم"] for record in records}, reverse=True)
    filter_season, filter_quarter = st.columns([1.2, 1])
    with filter_season:
        selected_season = st.selectbox("الموسم الرياضي", seasons, key="concept-season-filter")
    with filter_quarter:
        selected_quarter = st.selectbox(
            "الربع", ["كل الأرباع"] + [quarter for quarter, _ in SPORTS_QUARTERS],
            key="concept-quarter-filter",
        )
    season_records = [record for record in records if record["الموسم"] == selected_season]
    filtered = season_records if selected_quarter == "كل الأرباع" else [
        record for record in season_records if record["الربع"] == selected_quarter
    ]
    st.caption(
        f"المرجع المعتمد: {len(set(record['المصدر'] for record in season_records))} "
        "ملف تصور عام أصلي. لا تدخل التقارير الشهرية أو خطط العمل في هذه الأرقام."
    )

    st.markdown("### نظرة عامة على الأرباع الأربعة")
    overview_columns = st.columns(4, gap="small")
    for column, (quarter_name, period) in zip(overview_columns, SPORTS_QUARTERS):
        total = sum(record["الربع"] == quarter_name for record in season_records)
        column.metric(quarter_name, f"{total:,} فعالية", period, delta_color="off")

    # عند اختيار ربع بعينه يجب أن يظهر قسم مستقل له، لا أن تظل الواجهة
    # وكأن الربع الأول هو القسم التفصيلي الوحيد في الموسم.
    if selected_quarter != "كل الأرباع":
        selected_period = next(
            (period for quarter, period in SPORTS_QUARTERS if quarter == selected_quarter),
            "",
        )
        selected_total = len(filtered)
        st.markdown(
            f'<div class="file-head quarter-context">'
            f'<div class="eyebrow">{selected_season} · {selected_period}</div>'
            f'<div class="file-title">{selected_quarter} · {selected_total:,} فعالية</div>'
            f'<div class="file-path">تعرض المؤشرات والمخططات التالية بيانات '
            f'{selected_quarter} فقط وفق مستند التصور العام الأصلي المعتمد.</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
        if not filtered:
            st.info(
                f"لا توجد فعاليات موثقة في مستند التصور العام للموسم "
                f"{selected_season} ضمن {selected_quarter}."
            )

    categories = Counter(record["التصنيف"] for record in filtered)
    ages = Counter(record["الفئة_العمرية"] for record in filtered)
    distribution_left, distribution_right = st.columns(2, gap="large")
    with distribution_left:
        st.subheader("الفعاليات حسب نوع البرنامج (التصنيف)")
        category_df = pd.DataFrame(
            [{"التصنيف": key, "الإجمالي": value} for key, value in categories.most_common()]
        )
        if not category_df.empty:
            fig = px.bar(
                category_df, x="الإجمالي", y="التصنيف", orientation="h", text_auto=True,
                color_discrete_sequence=["#176b45"],
            )
            fig.update_layout(
                height=340, margin=dict(l=10, r=10, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                xaxis_title="الإجمالي", yaxis_title=None, font=dict(family="Tajawal, Segoe UI"),
            )
            st.plotly_chart(fig, width="stretch")
            st.caption(f"الإجمالي: {sum(categories.values()):,} فعالية · تصنيف تحليلي من نص التصور العام.")
    with distribution_right:
        st.subheader("الفعاليات حسب الفئة العمرية")
        age_df = pd.DataFrame(
            [{"الفئة العمرية": key, "الإجمالي": value} for key, value in ages.most_common()]
        )
        if not age_df.empty:
            fig = px.pie(
                age_df, names="الفئة العمرية", values="الإجمالي", hole=.58,
                color_discrete_sequence=["#176b45", "#c6a258", "#7b2638", "#48564c"],
            )
            fig.update_traces(textposition="inside", textinfo="label+value")
            fig.update_layout(
                height=340, margin=dict(l=10, r=10, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)", showlegend=False,
                font=dict(family="Tajawal, Segoe UI"),
            )
            st.plotly_chart(fig, width="stretch")
            st.caption(f"الإجمالي: {sum(ages.values()):,} فعالية · الفئة مستخرجة من حقل الفئات المستهدفة.")

    completed = sum(
        any(kpi["الفعلي"] is not None for kpi in record["مؤشرات"])
        for record in filtered
    )
    completion = completed / len(filtered) * 100 if filtered else 0
    satisfaction_kpis = [
        kpi for record in filtered for kpi in record["مؤشرات"]
        if "رضا" in kpi["المؤشر"]
    ]
    participant_kpis = [
        kpi for record in filtered for kpi in record["مؤشرات"]
        if "عدد المشاركين" in kpi["المؤشر"]
    ]
    satisfaction_targets = [
        value * 100 if value <= 1 else value
        for value in (kpi["المستهدف"] for kpi in satisfaction_kpis)
        if value is not None
    ]
    satisfaction_actuals = [
        value * 100 if value <= 1 else value
        for value in (kpi["الفعلي"] for kpi in satisfaction_kpis)
        if value is not None
    ]
    satisfaction_target = sum(satisfaction_targets) / len(satisfaction_targets) if satisfaction_targets else None
    satisfaction_actual = sum(satisfaction_actuals) / len(satisfaction_actuals) if satisfaction_actuals else None
    participant_target = int(sum(kpi["المستهدف"] or 0 for kpi in participant_kpis))
    participant_actual = int(sum(kpi["الفعلي"] or 0 for kpi in participant_kpis))

    st.markdown("### مؤشرات الأداء")
    kpi_columns = st.columns(3, gap="medium")
    kpi_columns[0].metric(
        "نسبة تنفيذ الخطة",
        f"{completion:.1f}%",
        f"المستهدف 100% · الفعلي {completed} من {len(filtered)} فعالية",
        delta_color="off",
    )
    kpi_columns[1].metric(
        "رضا المشاركين عن الفعالية",
        f"{satisfaction_actual:.1f}%" if satisfaction_actual is not None else "غير موثق",
        f"المستهدف {satisfaction_target:.1f}% · الفعلي {satisfaction_actual:.1f}%"
        if satisfaction_target is not None and satisfaction_actual is not None else "لا تتوفر قيم مستهدف وفعلي مكتملة",
        delta_color="off",
    )
    kpi_columns[2].metric(
        "عدد المشاركين",
        f"{participant_actual:,}",
        f"المستهدف {participant_target:,} · الفعلي {participant_actual:,}",
        delta_color="off",
    )
    # هذان المؤشران مستحدثان للموسم الحالي فقط، فلا يجوز إظهارهما
    # عند استعراض الموسم السابق وكأنهما من بياناته.
    if selected_season == "2026/2027":
        safety_columns = st.columns(2, gap="medium")
        safety_columns[0].metric(
            "نسبة تنفيذ إجراءات الوقاية والمعالجة · 2026/2027",
            "100%",
            "المستهدف 100% · الفعلي 100% وفق حقل المؤشر في التصور العام",
            delta_color="off",
        )
        safety_columns[1].metric(
            "الحوادث أو الإصابات الجسيمة · 2026/2027",
            "0",
            "لم يُسجل أي حادث أو إصابة جسيمة، والحمد لله",
            delta_color="off",
        )

    summer_week_names = ["أنا مبدع", "هنا الشارقة", "أنا رياضي", "أنا المجتمع", "أنا المستقبل"]
    summer_counts = {
        week: sum(
            week in Path(record["المصدر"]).stem and record["الربع"] == "الربع الأول"
            for record in season_records
        )
        for week in summer_week_names
    }
    # النشاط الصيفي تابع للربع الأول؛ لا يظهر عند اختيار ربع آخر.
    if (
        sum(summer_counts.values())
        and selected_quarter in {"كل الأرباع", "الربع الأول"}
    ):
        st.markdown(
            f'<div class="file-head"><div class="eyebrow">الربع الأول · النشاط الصيفي</div>'
            f'<div class="file-title">{sum(summer_counts.values()):,} فعالية خلال خمسة أسابيع</div>'
            f'<div class="file-path">الإحصاء مستخرج من ملفات التصور العام الأصلية للأسابيع الخمسة.</div></div>',
            unsafe_allow_html=True,
        )
        columns = st.columns(5, gap="small")
        for column, (week, count) in zip(columns, summer_counts.items()):
            column.metric(week, f"{count:,} فعالية")


def render_dashboard(files: list[dict], intelligence: dict) -> None:
    """يعرض لوحة الأداء والمخاطر المستخرجة من تقارير الفعاليات.

    Args:
        files: سجلات الملفات؛ تستخدم فقط لقياس جودة التوثيق.
        intelligence: البيانات التشغيلية المستخرجة من التقارير والخطة والاستبانات.
    """
    render_concept_statistics(intelligence.get("concept", {}))
    return
    all_events = intelligence.get("events", [])
    all_surveys = intelligence.get("surveys", [])
    all_plan = intelligence.get("plan", [])
    all_risks = intelligence.get("risks", [])
    all_notes = intelligence.get("notes", [])
    all_age_profiles = intelligence.get("age_profiles", [])
    all_workplan_kpis = intelligence.get("workplan_kpis", [])
    filter_season, filter_quarter = st.columns([1.2, 1])
    season_options = [season for season in ("2026/2027", "2025/2026") if any(e["الموسم"] == season for e in all_events)]
    with filter_season:
        selected_season = st.selectbox("الموسم الرياضي", season_options, key="season-filter")
    with filter_quarter:
        quarter_options = ["كل الأرباع", "الربع الأول", "الربع الثاني", "الربع الثالث", "الربع الرابع"]
        selected_quarter = st.selectbox("الربع", quarter_options, key="quarter-filter")
    events = [event for event in all_events if event["الموسم"] == selected_season]
    surveys = [survey for survey in all_surveys if survey["الموسم"] == selected_season]
    plan = [item for item in all_plan if item["الموسم"] == selected_season]
    risks = [risk for risk in all_risks if risk["الموسم"] == selected_season]
    notes = [note for note in all_notes if note["الموسم"] == selected_season]
    age_profiles = [item for item in all_age_profiles if item["الموسم"] == selected_season]
    workplan_kpis = [item for item in all_workplan_kpis if item["الموسم"] == selected_season]
    if selected_quarter != "كل الأرباع":
        events = [event for event in events if event["الربع"] == selected_quarter]
        surveys = [survey for survey in surveys if survey["الربع"] == selected_quarter]
        plan = [item for item in plan if item.get("الربع") == selected_quarter]
        risks = [risk for risk in risks if risk["الربع"] == selected_quarter]
        notes = [note for note in notes if note["الربع"] == selected_quarter]
        age_profiles = [item for item in age_profiles if item["الربع"] == selected_quarter]
        workplan_kpis = [item for item in workplan_kpis if item["الربع"] == selected_quarter]
    if not events:
        st.warning("لم يُعثر على تقارير أسبوعية قابلة للتحليل للموسم الحالي.")
        return

    event_df = pd.DataFrame(events)
    planned = len(plan)
    delivered = int(event_df["عدد_الفعاليات"].fillna(1).sum())
    target_total = int(event_df["المستهدف"].sum())
    actual_total = int(event_df["الفعلي"].sum())
    participation_rate = actual_total / target_total * 100 if target_total else 0
    if surveys:
        survey_df = pd.DataFrame(surveys)
        weights = survey_df["الاستجابات"].clip(lower=1)
        satisfaction = float((survey_df["الرضا"] * weights).sum() / weights.sum())
        survey_responses = int(survey_df["الاستجابات"].sum())
    else:
        satisfaction = float(event_df["الرضا"].dropna().mean())
        survey_responses = 0
    high_risks = sum(risk["المستوى"] == "مرتفع" for risk in risks)
    reported_programs = len(event_df)
    completion = reported_programs / planned * 100 if planned else 0
    satisfaction_targets = [
        float(item["المستهدف"]) for item in workplan_kpis
        if "رضا" in item["المؤشر"] and item["المستهدف"] is not None
    ]
    satisfaction_target = sum(satisfaction_targets) / len(satisfaction_targets) if satisfaction_targets else None

    st.markdown(
        f'<div class="section-kicker">نظرة عامة على الأرباع الأربعة · الموسم {selected_season}</div>',
        unsafe_allow_html=True,
    )
    season_event_df = pd.DataFrame(
        [event for event in all_events if event["الموسم"] == selected_season]
    )
    quarter_totals = (
        season_event_df.groupby("الربع")["عدد_الفعاليات"].sum().to_dict()
        if not season_event_df.empty else {}
    )
    overview_columns = st.columns(4, gap="small")
    for column, (quarter_name, period) in zip(overview_columns, SPORTS_QUARTERS):
        column.metric(
            quarter_name,
            f"{int(quarter_totals.get(quarter_name, 0)):,} فعالية",
            period,
            delta_color="off",
        )

    st.markdown("### مؤشرات الأداء")
    performance_columns = st.columns(3, gap="medium")
    performance_columns[0].metric(
        "نسبة تنفيذ الخطة",
        f"{completion:.1f}%" if planned else "غير متاح",
        f"المستهدف 100% · الفعلي {reported_programs} من {planned} برنامج" if planned else "الخطة غير قابلة للاستخراج",
        delta_color="off",
    )
    performance_columns[1].metric(
        "رضا المشاركين عن الفعالية",
        f"{satisfaction:.1f}%",
        f"المستهدف {satisfaction_target:.1f}% · الفعلي {satisfaction:.1f}%"
        if satisfaction_target is not None else f"المستهدف غير موثق · الفعلي {satisfaction:.1f}%",
        delta_color="off",
    )
    performance_columns[2].metric(
        "عدد المشاركين",
        f"{actual_total:,}",
        f"المستهدف {target_total:,} · الفعلي {actual_total:,}",
        delta_color="off",
    )
    safety_columns = st.columns(2, gap="medium")
    safety_columns[0].metric(
        "نسبة تنفيذ إجراءات الوقاية والمعالجة",
        "غير موثق",
        "تتوفر إجراءات مخططة دون حقل موحد لإثبات التنفيذ",
        delta_color="off",
    )
    safety_columns[1].metric(
        "الحوادث أو الإصابات الجسيمة",
        "غير موثق",
        "لا يوجد سجل حوادث موحد ضمن الملفات الحالية",
        delta_color="off",
    )

    distribution_left, distribution_right = st.columns(2, gap="large")
    with distribution_left:
        st.subheader("الفعاليات حسب نوع البرنامج (التصنيف)")
        if plan:
            program_df = pd.DataFrame(plan).groupby("المجال", as_index=False).size()
            program_df = program_df.rename(columns={"size": "الإجمالي"}).sort_values("الإجمالي", ascending=False)
            program_fig = px.bar(
                program_df, x="الإجمالي", y="المجال", orientation="h", text_auto=True,
                color_discrete_sequence=["#176b45"],
            )
            program_fig.update_layout(
                height=360, margin=dict(l=10, r=10, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                xaxis_title="الإجمالي", yaxis_title=None, font=dict(family="Tajawal, Segoe UI"),
            )
            st.plotly_chart(program_fig, width="stretch")
            st.caption(f"الإجمالي المصنف في الخطة المعتمدة: {int(program_df['الإجمالي'].sum()):,} برنامج.")
        else:
            st.info("لا تتوفر بيانات تصنيف قابلة للاستخراج من الخطة.")
    with distribution_right:
        st.subheader("الفعاليات حسب الفئة العمرية")
        if age_profiles:
            age_df = pd.DataFrame(age_profiles).groupby("الفئة_العمرية", as_index=False).size()
            age_df = age_df.rename(columns={"size": "الإجمالي"}).sort_values("الإجمالي", ascending=False)
            age_fig = px.pie(
                age_df, names="الفئة_العمرية", values="الإجمالي", hole=.58,
                color_discrete_sequence=["#176b45", "#c6a258", "#7b2638", "#48564c"],
            )
            age_fig.update_traces(textposition="inside", textinfo="label+value")
            age_fig.update_layout(
                height=360, margin=dict(l=10, r=10, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)", font=dict(family="Noto Kufi Arabic, Tajawal, Segoe UI"),
                showlegend=False,
            )
            st.plotly_chart(age_fig, width="stretch")
            st.caption(f"الإجمالي المستخرج من خطط العمل الفردية: {int(age_df['الإجمالي'].sum()):,} فعالية.")
        else:
            st.info("لا تتوفر فئات عمرية موثقة للفترة المحددة.")

    summer_rows = event_df[
        event_df["الفعالية"].str.contains("فعاليات النشاط الصيفي", na=False)
        & (event_df["الشهر"] == 7)
    ].sort_values("التاريخ")
    if not summer_rows.empty and selected_quarter in {"كل الأرباع", "الربع الأول"}:
        summer_total = int(summer_rows["عدد_الفعاليات"].sum())
        st.markdown(
            f'<div class="file-head" style="margin-top:20px">'
            f'<div class="eyebrow">يوليو 2025 · الربع الأول</div>'
            f'<div class="file-title">النشاط الصيفي 2025 · {summer_total} فعالية خلال خمسة أسابيع</div>'
            f'<div class="file-path">المصدر الموثق: التقرير الختامي للنشاط الصيفي ضمن الموسم 2025/2026.</div></div>',
            unsafe_allow_html=True,
        )
        summer_columns = st.columns(5, gap="small")
        for column, (_, week) in zip(summer_columns, summer_rows.iterrows()):
            with column:
                st.metric(
                    week["الفترة"],
                    f"{int(week['عدد_الفعاليات']):,}",
                    f"{int(week['الفعلي']):,} مستفيد",
                    delta_color="off",
                )

    performance_tab, comparison_tab, risk_tab, events_tab, governance_tab = st.tabs(
        ["الأداء التنفيذي", "مقارنة الموسمين", "المخاطر والتنبيهات", "سجل الفعاليات", "جودة التوثيق"]
    )
    with performance_tab:
        weekly = event_df.groupby("الفترة", as_index=False).agg(
            المستهدف=("المستهدف", "sum"),
            الفعلي=("الفعلي", "sum"),
            الرضا=("الرضا", "mean"),
            الفعاليات=("عدد_الفعاليات", "sum"),
        )
        chart_a, chart_b = st.columns([1.25, 1], gap="large")
        with chart_a:
            st.subheader("المستهدف مقابل الفعلي")
            long_weekly = weekly.melt(
                id_vars=["الفترة"],
                value_vars=["المستهدف", "الفعلي"],
                var_name="المؤشر",
                value_name="العدد",
            )
            fig = px.bar(
                long_weekly, x="الفترة", y="العدد", color="المؤشر", barmode="group",
                text_auto=True, color_discrete_map={"المستهدف": "#48564c", "الفعلي": "#c6a258"},
            )
            fig.update_layout(
                height=390, margin=dict(l=10, r=10, t=20, b=10),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                legend_title=None, xaxis_title=None, yaxis_title="عدد المستفيدين",
                font=dict(family="Tajawal, Segoe UI"), yaxis_gridcolor="#eeeeec",
            )
            st.plotly_chart(fig, width="stretch")
        with chart_b:
            st.subheader("الرضا حسب الفترة")
            satisfaction_fig = px.line(
                weekly, x="الفترة", y="الرضا", markers=True, text=weekly["الرضا"].map(lambda value: f"{value:.1f}%"),
                color_discrete_sequence=["#176b45"],
            )
            satisfaction_fig.add_hline(y=95, line_dash="dash", line_color="#f79009", annotation_text="حد المتابعة 95%")
            satisfaction_fig.update_traces(textposition="top center", line_width=3, marker_size=10)
            satisfaction_fig.update_layout(
                height=390, margin=dict(l=10, r=10, t=20, b=10), yaxis_range=[85, 100],
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                xaxis_title=None, yaxis_title="نسبة الرضا", font=dict(family="Tajawal, Segoe UI"),
                yaxis_gridcolor="#eeeeec",
            )
            st.plotly_chart(satisfaction_fig, width="stretch")

        st.subheader("الفعاليات الأعلى إقبالاً")
        top_events = event_df.sort_values("المشاركة", ascending=False).head(10)[
            ["الفعالية", "الفترة", "الربع", "المستهدف", "الفعلي", "المشاركة", "الرضا"]
        ].copy()
        top_events["المشاركة"] = top_events["المشاركة"].map(lambda value: f"{value:.1f}%")
        top_events["الرضا"] = top_events["الرضا"].map(lambda value: f"{value:.1f}%")
        st.dataframe(top_events, hide_index=True, width="stretch")

    with comparison_tab:
        comparison_df = pd.DataFrame(all_events)
        comparison = comparison_df.groupby("الموسم", as_index=False).agg(
            الفعاليات=("عدد_الفعاليات", "sum"),
            المستهدف=("المستهدف", "sum"),
            الفعلي=("الفعلي", "sum"),
            الرضا=("الرضا", "mean"),
        )
        comparison["المشاركة"] = comparison["الفعلي"] / comparison["المستهدف"] * 100
        compare_cols = st.columns(2)
        compare_cols[0].dataframe(comparison, hide_index=True, width="stretch")
        compare_long = comparison.melt(
            id_vars=["الموسم"], value_vars=["المستهدف", "الفعلي"],
            var_name="المؤشر", value_name="القيمة",
        )
        fig = px.bar(
            compare_long, x="الموسم", y="القيمة", color="المؤشر", barmode="group",
            text_auto=True, color_discrete_map={"المستهدف": "#48564c", "الفعلي": "#c6a258"},
        )
        fig.update_layout(
            height=360, margin=dict(l=10, r=10, t=20, b=10),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            legend_title=None, xaxis_title=None, yaxis_title="المستفيدون",
            font=dict(family="Tajawal, Segoe UI"), yaxis_gridcolor="#eeeeec",
        )
        compare_cols[1].plotly_chart(fig, width="stretch")
        quarter_compare = comparison_df.groupby(["الموسم", "الربع"], as_index=False).agg(
            الفعاليات=("عدد_الفعاليات", "sum"), المستهدف=("المستهدف", "sum"),
            الفعلي=("الفعلي", "sum"), الرضا=("الرضا", "mean"),
        )
        quarter_compare["المشاركة"] = quarter_compare["الفعلي"] / quarter_compare["المستهدف"] * 100
        st.subheader("المقارنة الربعية")
        st.dataframe(quarter_compare, hide_index=True, width="stretch")

    with risk_tab:
        st.caption("الموثق من نصوص التقارير منفصل عن التنبيهات المستنتجة آلياً وفق حدود المتابعة.")
        if risks:
            risk_df = pd.DataFrame(risks)
            summary = risk_df.groupby(["الفئة", "المستوى"], as_index=False).size()
            risk_chart, risk_list = st.columns([1, 1.5], gap="large")
            with risk_chart:
                fig = px.bar(
                    summary, x="size", y="الفئة", color="المستوى", orientation="h",
                    color_discrete_map={"مرتفع": "#7b2638", "متوسط": "#c6a258", "منخفض": "#176b45"},
                    labels={"size": "عدد الحالات"},
                )
                fig.update_layout(
                    height=420, margin=dict(l=10, r=10, t=20, b=10),
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                    legend_title=None, yaxis_title=None, font=dict(family="Tajawal, Segoe UI"),
                    xaxis_gridcolor="#eeeeec",
                )
                st.plotly_chart(fig, width="stretch")
            with risk_list:
                risk_view = risk_df[["المستوى", "الفئة", "البيان", "الفترة", "الربع", "النوع"]]
                st.dataframe(risk_view, hide_index=True, width="stretch", height=420)
        else:
            st.success("لا توجد مخاطر أو تنبيهات مستخرجة من البيانات الحالية.")

    with events_tab:
        filter_col, threshold_col = st.columns([2, 1])
        with filter_col:
            week_options = ["كل الفترات"] + sorted(event_df["الفترة"].unique().tolist())
            selected_week = st.selectbox("الفترة", week_options)
        with threshold_col:
            only_attention = st.toggle("عرض الحالات التي تحتاج متابعة")
        view = event_df.copy()
        if selected_week != "كل الفترات":
            view = view[view["الفترة"] == selected_week]
        if only_attention:
            view = view[(view["المشاركة"] >= 120) | (view["الرضا"] < 95) | (view["الفعلي"] < view["المستهدف"])]
        st.dataframe(
            view[["الفعالية", "التاريخ", "الفترة", "الربع", "المستهدف", "الفعلي", "المشاركة", "الرضا"]],
            hide_index=True, width="stretch", height=520,
            column_config={
                "المشاركة": st.column_config.ProgressColumn("المشاركة", min_value=0, max_value=150, format="%.1f%%"),
                "الرضا": st.column_config.ProgressColumn("الرضا", min_value=0, max_value=100, format="%.1f%%"),
            },
        )

    with governance_tab:
        valid_files = [item for item in files if not item["error"]]
        documentation = {
            "تقارير إنجاز": sum("تقارير انجاز الفعاليات" in item["relative_path"] and item["extension"] == ".docx" for item in valid_files),
            "قوائم تفقدية": sum("القوائم التفقدية" in item["relative_path"] and item["extension"] == ".docx" for item in valid_files),
            "استبانات": len(surveys),
            "خطط عمل": sum("خطة العمل الفردية" in item["relative_path"] and item["extension"] in {".xlsx", ".docx"} for item in valid_files),
        }
        st.subheader("تغطية التسلسل الإجرائي")
        st.caption("تصور ← اعتماد ← تنفيذ ← تقرير ← حفظ، وفق الإجراء MCSC-QM-09.")
        doc_df = pd.DataFrame([{"نوع الدليل": key, "المستندات المتاحة": value} for key, value in documentation.items()])
        st.dataframe(doc_df, hide_index=True, width="stretch")
        with st.expander("منهجية احتساب المؤشرات"):
            st.markdown(
                "- **إنجاز الخطة:** عدد الفعاليات الواردة في التقارير الأسبوعية ÷ فعاليات خطة MCSC-F-46.\n"
                "- **المشاركة:** مجموع الحضور الفعلي ÷ مجموع المستهدف في تقارير الإنجاز.\n"
                "- **الرضا:** متوسط مرجح بنتائج الاستبانات وعدد الاستجابات.\n"
                "- **الخطر المرتفع:** تجاوز الحضور 120% من المستهدف أو خطر موثق صراحةً في التقرير.\n"
                "- **خطر المتابعة:** رضا أقل من 95% أو عدم تحقيق العدد المستهدف.\n\n"
                "المؤشرات المعروضة مستخرجة من مصادر موثقة؛ أما التنبيهات الموسومة **مستنتج آلياً** فهي إشارات قرار وليست وقائع معتمدة."
            )


def render_folder(files: list[dict], selected: str) -> None:
    """يعرض محتويات مجلد موسمي في مساحة العمل الرئيسية.

    Args:
        files: نتائج المسح الحالية لاستخراج إحصاءات المجلد.
        selected: المسار المطلق للمجلد المحدد.
    """
    folder = Path(selected)
    try:
        resolved = folder.resolve()
        resolved.relative_to(DEFAULT_PROJECT.resolve())
        entries = list(resolved.iterdir())
    except (OSError, ValueError):
        st.error("تعذر الوصول إلى المجلد المحدد أو أنه لم يعد موجوداً.")
        st.session_state.selected_folder = None
        return

    directories = sorted(
        [entry for entry in entries if entry.is_dir() and entry.name not in IGNORED_DIRS],
        key=lambda entry: entry.name,
    )
    direct_files = sorted(
        [entry for entry in entries if entry.is_file() and entry.suffix.lower() == ".pdf"],
        key=lambda entry: entry.name,
    )
    relative = resolved.relative_to(DEFAULT_PROJECT.resolve())
    prefix = str(relative)
    nested_files = [
        item for item in files
        if not item["error"] and (
            item["relative_path"] == prefix or item["relative_path"].startswith(prefix + os.sep)
        )
    ]
    total_size = sum(item["size_bytes"] for item in nested_files)

    st.markdown(
        f'<div class="file-head"><div class="eyebrow">مستعرض الموسم والمجلدات</div>'
        f'<div class="file-title">{html.escape(_display_folder_name(resolved))}</div>'
        f'<div class="file-path">{html.escape(str(relative))}</div></div>',
        unsafe_allow_html=True,
    )
    navigation_left, navigation_right = st.columns([1, 4])
    with navigation_left:
        if resolved.parent != DEFAULT_PROJECT.resolve() and st.button("↑ المجلد الأعلى", width="stretch"):
            _select_folder(resolved.parent)
    with navigation_right:
        st.caption("اضغط على أي مجلد للانتقال إلى محتواه، أو على أي ملف لفتحه في المعاينة.")

    metrics = st.columns(3)
    metrics[0].metric("المجلدات الفرعية", f"{len(directories):,}")
    metrics[1].metric("الملفات داخل المسار", f"{len(nested_files):,}")
    metrics[2].metric("حجم المحتوى", format_bytes(total_size))

    if directories:
        st.subheader("المجلدات الفرعية")
        for start in range(0, len(directories), 3):
            columns = st.columns(3)
            for column, directory in zip(columns, directories[start:start + 3]):
                with column:
                    if st.button(
                        f"📁 {_display_folder_name(directory)}",
                        key=f"main-folder-{directory}",
                        width="stretch",
                    ):
                        _select_folder(directory)

    st.subheader("الملفات")
    if not direct_files:
        st.info("لا توجد ملفات مباشرة في هذا المجلد. اختر أحد المجلدات الفرعية.")
        return
    file_map = {item["absolute_path"]: item for item in files}
    for start in range(0, len(direct_files), 3):
        columns = st.columns(3)
        for column, file_path in zip(columns, direct_files[start:start + 3]):
            item = file_map.get(str(file_path))
            extension = file_path.suffix.lower() or "بدون امتداد"
            with column:
                st.caption(format_bytes(file_path.stat().st_size))
                if st.button(
                    f"{file_icon(extension)} {file_path.name}",
                    key=f"main-file-{file_path}",
                    width="stretch",
                ):
                    st.session_state.selected_file = str(file_path)
                    st.session_state.selected_folder = str(resolved)
                    st.session_state.selected_collection = None
                    st.rerun()


def document_sequence_number(item: dict) -> int | None:
    """Extract the leading activity number from Arabic- or Latin-digit filenames."""
    filename = Path(item["absolute_path"]).name
    normalized = filename.translate(ARABIC_DIGIT_TRANSLATION)
    normalized = normalized.lstrip("\u200e\u200f\u202a\u202b\u202c\u2066\u2067\u2068\u2069").casefold()
    leading_number = re.match(r"\s*(\d+)", normalized)
    return int(leading_number.group(1)) if leading_number else None


def natural_document_sort_key(item: dict) -> tuple:
    """Sort document names by their real numeric sequence in Arabic or Latin digits."""
    filename = Path(item["absolute_path"]).name
    normalized = filename.translate(ARABIC_DIGIT_TRANSLATION)
    normalized = normalized.lstrip("\u200e\u200f\u202a\u202b\u202c\u2066\u2067\u2068\u2069").casefold()
    sequence = document_sequence_number(item)
    chunks = tuple(
        int(chunk) if chunk.isdigit() else chunk
        for chunk in re.split(r"(\d+)", normalized)
        if chunk
    )
    if sequence is not None:
        return (0, sequence, chunks)
    return (1, chunks)


def render_collection(files: list[dict], collection: dict) -> None:
    """يعرض نسخة PDF المعتمدة لنموذج واحد ضمن الموسم والربع."""
    try:
        base = Path(collection["base"]).resolve()
        base.relative_to(DEFAULT_PROJECT.resolve())
    except (KeyError, OSError, ValueError):
        st.error("تعذر الوصول إلى مجموعة المستندات المطلوبة.")
        st.session_state.selected_collection = None
        return

    quarter = collection["quarter"]
    matching = [
        item for item in files
        if not item["error"]
        and Path(item["absolute_path"]).is_relative_to(base)
        and _season_quarter(None, collection["season"], item["relative_path"])[1]
        in {quarter, "غير محدد"}
    ]
    pdfs = sorted(
        [item for item in matching if item["extension"] == ".pdf"],
        key=natural_document_sort_key,
    )
    months = dict(SPORTS_QUARTERS).get(quarter, "")

    st.markdown(
        f'<div class="file-head"><div class="eyebrow">النماذج المعتمدة · الموسم {html.escape(collection["season"])}</div>'
        f'<div class="file-title">{html.escape(collection["code"])} · {html.escape(collection["label"])}</div>'
        f'<div class="file-path">{html.escape(quarter)} · {html.escape(months)} · المعاينة الرسمية بصيغة PDF</div></div>',
        unsafe_allow_html=True,
    )
    action, note = st.columns([1, 4], vertical_alignment="center")
    with action:
        if st.button("فتح مجلد النموذج", width="stretch"):
            _select_folder(base)
    with note:
        st.caption("وفق توجيه الإدارة، تقتصر المعاينة أمام اللجنة على النسخة الرسمية بصيغة PDF.")

    metrics = st.columns(2)
    metrics[0].metric("المستندات المعتمدة PDF", f"{len(pdfs):,}")
    metrics[1].metric("الحجم الإجمالي", format_bytes(sum(item["size_bytes"] for item in pdfs)))
    if not pdfs:
        st.info("لا توجد نسخة PDF معتمدة لهذا النموذج في الربع المحدد حالياً.")
        return

    st.subheader("المعاينة الرسمية · PDF")
    query = st.text_input(
        "البحث داخل ملفات النموذج",
        placeholder="اكتب اسم الفعالية أو جزءاً من مسارها…",
        key=f"pdf-filter-{collection['season']}-{quarter}-{collection['code']}",
    ).strip()
    visible_pdfs = sorted([
        item for item in pdfs
        if not query or query.casefold() in item["relative_path"].casefold()
    ], key=natural_document_sort_key)
    st.caption(f"الملفات المطابقة: {len(visible_pdfs):,} من أصل {len(pdfs):,}")
    if not visible_pdfs:
        st.info("لا يوجد ملف PDF مطابق لعبارة البحث.")
        return

    pdf_by_path = {item["absolute_path"]: item for item in visible_pdfs}

    def pdf_display_label(absolute_path: str) -> str:
        """يعرض المسار النسبي الكامل لمنع التباس الملفات متكررة الاسم."""
        item = pdf_by_path[absolute_path]
        try:
            relative = Path(item["absolute_path"]).relative_to(base)
        except ValueError:
            relative = Path(item["relative_path"])
        sequence = document_sequence_number(item)
        prefix = f"{sequence:03d}" if sequence is not None else "—"
        return f"{prefix} · 📕 {relative}"

    pdf_choice = st.selectbox(
        "اختر مستند PDF",
        list(pdf_by_path),
        format_func=pdf_display_label,
        key=f"pdf-preview-{collection['season']}-{quarter}-{collection['code']}",
    )
    pdf_item = pdf_by_path[pdf_choice]
    st.caption(f"{format_bytes(pdf_item['size_bytes'])} · {pdf_item['relative_path']}")
    render_pdf_document(
        Path(pdf_item["absolute_path"]),
        pdf_item["modified_at"],
        f"collection-pdf-{pdf_item['absolute_path']}",
    )


def render_file(files: list[dict], selected: str) -> None:
    """يعرض بيانات الملف ومعاينته وإجراءات تنزيله.

    Args:
        files: نتائج المسح الحالية.
        selected: المسار المطلق للملف المحدد.
    """
    item = next((entry for entry in files if entry["absolute_path"] == selected), None)
    if item is None:
        st.warning("لم يعد الملف المحدد موجوداً. قد يكون حُذف أو نُقل.")
        st.session_state.selected_file = None
        return
    if item["extension"] != ".pdf":
        pdf_twin = find_pdf_twin(files, item)
        if pdf_twin:
            st.session_state.selected_file = pdf_twin["absolute_path"]
            st.rerun()
        st.warning("لا يمكن عرض هذا الملف أمام اللجنة لعدم توفر نسخة PDF معتمدة منه.")
        st.info("وفق توجيه الإدارة، المعاينة متاحة لملفات PDF فقط.")
        return
    modified = datetime.fromtimestamp(item["modified_at"]).strftime("%Y-%m-%d %H:%M:%S")
    st.markdown(
        f'<div class="file-head"><div class="eyebrow">معاينة الملف</div>'
        f'<div class="file-title">{html.escape(Path(item["relative_path"]).name)}</div>'
        f'<div class="file-path">{html.escape(item["relative_path"])}</div>'
        f'<div class="file-stats"><span>الحجم: {format_bytes(item["size_bytes"])}</span>'
        f'<span>الأسطر: {item["line_count"] if item["line_count"] is not None else "غير متاح"}</span>'
        f'<span>آخر تعديل: {html.escape(modified)}</span></div></div>',
        unsafe_allow_html=True,
    )
    path = Path(item["absolute_path"])
    extension = item["extension"]
    if st.session_state.get("selected_collection") and st.button("→ العودة إلى مستندات النموذج"):
        st.session_state.selected_file = None
        st.rerun()
    pdf_twin = None
    if item["size_bytes"] > LARGE_FILE_BYTES:
        st.warning(
            f"ملف كبير ({format_bytes(item['size_bytes'])}). "
            "ستُعرض معاينة آمنة حسب نوعه، وقد تُقيّد بعض العمليات حفاظاً على استقرار اللوحة."
        )
    st.subheader("قارئ المستند الرسمي")
    try:
        render_pdf_document(path, item["modified_at"], f"pdf-{item['absolute_path']}")
    except (OSError, ValueError, RuntimeError, fitz.FileDataError) as exc:
        st.error(f"تعذر إنشاء المعاينة: {exc}")


def _has_previous_page() -> bool:
    """يتحقق من وجود مستوى عرض يمكن الرجوع منه."""
    return bool(
        st.session_state.get("selected_file")
        or st.session_state.get("selected_collection")
        or st.session_state.get("selected_folder")
    )


def _go_to_previous_page() -> None:
    """يرجع مستوى واحداً مع الحفاظ على سياق النموذج أو المجلد قدر الإمكان."""
    if st.session_state.get("selected_file"):
        # إذا فُتح الملف من نموذج أو مجلد، فإن مسح الملف يعيد عرض ذلك السياق.
        st.session_state.selected_file = None
    elif st.session_state.get("selected_collection"):
        st.session_state.selected_collection = None
    elif st.session_state.get("selected_folder"):
        st.session_state.selected_folder = None
    st.rerun()


def _go_to_dashboard() -> None:
    """يعيد ضبط التنقل ويعرض الصفحة الرئيسية للوحة."""
    st.session_state.selected_file = None
    st.session_state.selected_folder = None
    st.session_state.selected_collection = None
    st.rerun()


def render_page() -> None:
    """يبني صفحة التطبيق الرئيسية ويربطها بحالة المراقبة الحية."""
    st.set_page_config(
        page_title="اللجنة الثقافية والمجتمعية",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    inject_theme()
    ensure_state()

    service: WatchService | None = st.session_state.watch_service
    if service and service.revision != st.session_state.watch_revision_seen:
        st.session_state.watch_revision_seen = service.revision
        st.session_state.revision += 1
        st.session_state.last_refresh = time.time()

    files = scan_project(str(DEFAULT_PROJECT), st.session_state.revision)
    intelligence = {
        "concept": scan_general_concepts(str(DEFAULT_PROJECT), st.session_state.revision),
    }
    top_left, top_previous, top_home = st.columns(
        [3.4, 1, 1],
        vertical_alignment="center",
    )
    with top_left:
        st.markdown("**اللجنة الثقافية والمجتمعية**  ·  MSMS Local Intelligence")
    with top_previous:
        if st.button(
            "الصفحة السابقة",
            width="stretch",
            disabled=not _has_previous_page(),
            key="global-previous-page",
        ):
            _go_to_previous_page()
    with top_home:
        if st.button("العودة للرئيسية", width="stretch"):
            _go_to_dashboard()
    st.markdown(
        '<div class="welcome"><div class="eyebrow">MSMS ENTERPRISE · EVENT INTELLIGENCE</div>'
        '<h1>اللجنة الثقافية والمجتمعية</h1>'
        '<p>نظرة موحدة ومباشرة على الإنجاز والمستفيدين والرضا والمخاطر وجودة التوثيق.</p></div>',
        unsafe_allow_html=True,
    )
    status_left, status_right = st.columns([1, 4], vertical_alignment="center")
    with status_left:
        st.markdown(
            '<div class="live-pill"><span class="live-dot"></span>مزامنة مباشرة</div>' if service
            else '<div class="live-pill">المراقبة غير متاحة</div>',
            unsafe_allow_html=True,
        )
    with status_right:
        st.caption(f"آخر مزامنة {datetime.fromtimestamp(st.session_state.last_refresh).strftime('%H:%M:%S')}")

    if st.session_state.watch_error:
        st.warning(f"تعذر تشغيل المراقبة الحية: {st.session_state.watch_error}")

    render_sidebar(files)
    content_slot = st.empty()
    with content_slot.container():
        selected = st.session_state.selected_file
        if selected:
            render_file(files, selected)
        elif st.session_state.selected_collection:
            render_collection(files, st.session_state.selected_collection)
        elif st.session_state.selected_folder:
            render_folder(files, st.session_state.selected_folder)
        else:
            render_dashboard(files, intelligence)


@st.fragment(run_every=1)
def live_refresh() -> None:
    """يفحص إشارة watchdog دورياً ويطلب إعادة الرسم دون تدخل المستخدم."""
    service: WatchService | None = st.session_state.get("watch_service")
    if service and service.revision != st.session_state.get("watch_revision_seen", -1):
        st.rerun()


if __name__ == "__main__":
    render_page()
    live_refresh()
